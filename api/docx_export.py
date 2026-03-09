# api/docx_export.py

import json
import os
import re
import io
import zipfile
import tempfile
import logging
import subprocess
from pathlib import Path
from io import BytesIO

import requests
from lxml import etree

from django.conf import settings
from django.http import JsonResponse, HttpResponse
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.http import require_http_methods
from django.shortcuts import get_object_or_404
from django.template.loader import render_to_string

from docx import Document
from docx.shared import Mm, Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .models import Curso
from .decorators import require_admin_token

from api.horario_legend import (
    build_horario_header_from_db,
    build_modules_legend_from_db,
    _add_legend_nonlective,
    _add_modules_legend_blocks,
)

logger = logging.getLogger(__name__)

LIBREOFFICE_BIN = getattr(settings, "LIBREOFFICE_PATH", "/usr/bin/soffice")
MEATZE_DOCX_LOGOS = getattr(settings, "MEATZE_DOCX_LOGOS", {})

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS = {"w": W_NS}


# ===== DOCX markers (module-level) =====
COURSE_INFO_MARKERS = [
    "Nombre curso:",
    "Código curso:",
    "Año académico:",
    "Fechas impartición:",
    "Horario:",
    "Entidad:",
    "Tipo formación:",
]

import zipfile
from lxml import etree
from pathlib import Path

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS = {"w": W_NS}

def _patch_docx_spacing_between_anchors(
    docx_path: Path,
    start_anchor: str,
    end_anchor: str | None = None,
    before="0",
    after="0",
    line=240,
    line_rule="auto",
    course_markers=None,   # ✅ NEW
):
    if course_markers is None:
        course_markers = [
            "Nombre curso:",
            "Código curso:",
            "Año académico:",
            "Fechas impartición:",
            "Horario:",
            "Entidad:",
            "Tipo formación:",
        ]

    def q(tag): return f"{{{W_NS}}}{tag}"

    def ensure(parent, tag):
        el = parent.find(f"w:{tag}", NS)
        if el is None:
            el = etree.Element(q(tag))
            parent.append(el)
        return el

    def node_text(node):
        parts = node.xpath(".//w:t/text()", namespaces=NS)
        return " ".join(" ".join(parts).split()).strip()

    def is_course_table(tbl):
        found = 0
        for tr in tbl.xpath("./w:tr", namespaces=NS):
            tcs = tr.xpath("./w:tc", namespaces=NS)
            if not tcs:
                continue
            left = node_text(tcs[0])
            if any(left.startswith(m) for m in course_markers):
                found += 1
        return found >= 3

    with zipfile.ZipFile(docx_path, "r") as zin:
        files = {name: zin.read(name) for name in zin.namelist()}

    xml = files.get("word/document.xml")
    if not xml:
        return 0

    root = etree.fromstring(xml)
    ps = root.xpath(".//w:p", namespaces=NS)

    started = False
    start_l = (start_anchor or "").strip().lower()
    end_l   = (end_anchor or "").strip().lower() if end_anchor else None

    patched = 0

    for p in ps:
        txt = node_text(p).lower()

        if not started:
            if start_l and start_l in txt:
                started = True
            continue

        if end_l and end_l in txt:
            break

        # ✅ если абзац внутри таблицы — скипаем ТОЛЬКО таблицу курса
        tbl_anc = p.xpath("ancestor::w:tbl[1]", namespaces=NS)
        if tbl_anc:
            if is_course_table(tbl_anc[0]):
                continue  # курс-таблица не трогаем, чтобы не ломать

        pPr = p.find("w:pPr", NS)
        if pPr is None:
            pPr = etree.Element(q("pPr"))
            p.insert(0, pPr)

        spacing = ensure(pPr, "spacing")
        spacing.set(q("before"), str(before))
        spacing.set(q("after"), str(after))
        spacing.set(q("line"), str(line))
        spacing.set(q("lineRule"), str(line_rule))

        cs = pPr.find("w:contextualSpacing", NS)
        if cs is not None:
            pPr.remove(cs)

        patched += 1

    files["word/document.xml"] = etree.tostring(
        root, xml_declaration=True, encoding="UTF-8", standalone="yes"
    )

    with zipfile.ZipFile(docx_path, "w", compression=zipfile.ZIP_DEFLATED) as zout:
        for n, data in files.items():
            zout.writestr(n, data)

    return patched

def _remove_empty_paragraphs_between_anchors(
    docx_path: Path,
    start_anchor: str,
    end_anchor: str | None = None,
    skip_tables=True,
):
    def node_text(node):
        parts = node.xpath(".//w:t/text()", namespaces=NS)
        return " ".join(" ".join(parts).split()).strip()

    with zipfile.ZipFile(docx_path, "r") as zin:
        files = {name: zin.read(name) for name in zin.namelist()}

    xml = files.get("word/document.xml")
    if not xml:
        return 0

    root = etree.fromstring(xml)
    ps = root.xpath(".//w:p", namespaces=NS)

    start_l = (start_anchor or "").strip().lower()
    end_l   = (end_anchor or "").strip().lower() if end_anchor else None

    started = False
    removed = 0

    # Важно: удалять нужно с конца, поэтому сначала соберём кандидатов
    to_remove = []

    for p in ps:
        txt = node_text(p).lower()

        if not started:
            if start_l and start_l in txt:
                started = True
            continue

        if end_l and end_l in txt:
            break

        if skip_tables and p.xpath("ancestor::w:tbl", namespaces=NS):
            continue

        # если в абзаце есть рисунки/объекты — не трогаем
        if p.xpath(".//w:drawing | .//w:pict | .//w:object", namespaces=NS):
            continue

        # пустой текст (включая неразрывные пробелы, которые Word часто кладёт)
        t = node_text(p).replace("\u00a0", "").strip()
        if t == "":
            to_remove.append(p)

    for p in reversed(to_remove):
        parent = p.getparent()
        if parent is not None:
            parent.remove(p)
            removed += 1

    files["word/document.xml"] = etree.tostring(
        root, xml_declaration=True, encoding="UTF-8", standalone="yes"
    )

    with zipfile.ZipFile(docx_path, "w", compression=zipfile.ZIP_DEFLATED) as zout:
        for n, data in files.items():
            zout.writestr(n, data)

    return removed

def _patch_docx_spacing_between_anchors(
    docx_path: Path,
    start_anchor: str,
    end_anchor: str | None = None,
    before="0",
    after="0",
    line=240,
    line_rule="auto",
    course_markers=None,   # ✅ NEW
):
    if course_markers is None:
        course_markers = [
            "Nombre curso:",
            "Código curso:",
            "Año académico:",
            "Fechas impartición:",
            "Horario:",
            "Entidad:",
            "Tipo formación:",
        ]

    def q(tag): return f"{{{W_NS}}}{tag}"

    def ensure(parent, tag):
        el = parent.find(f"w:{tag}", NS)
        if el is None:
            el = etree.Element(q(tag))
            parent.append(el)
        return el

    def node_text(node):
        parts = node.xpath(".//w:t/text()", namespaces=NS)
        return " ".join(" ".join(parts).split()).strip()

    def is_course_table(tbl):
        found = 0
        for tr in tbl.xpath("./w:tr", namespaces=NS):
            tcs = tr.xpath("./w:tc", namespaces=NS)
            if not tcs:
                continue
            left = node_text(tcs[0])
            if any(left.startswith(m) for m in course_markers):
                found += 1
        return found >= 3

    with zipfile.ZipFile(docx_path, "r") as zin:
        files = {name: zin.read(name) for name in zin.namelist()}

    xml = files.get("word/document.xml")
    if not xml:
        return 0

    root = etree.fromstring(xml)
    ps = root.xpath(".//w:p", namespaces=NS)

    started = False
    start_l = (start_anchor or "").strip().lower()
    end_l   = (end_anchor or "").strip().lower() if end_anchor else None

    patched = 0

    for p in ps:
        txt = node_text(p).lower()

        if not started:
            if start_l and start_l in txt:
                started = True
            continue

        if end_l and end_l in txt:
            break

        # ✅ если абзац внутри таблицы — скипаем ТОЛЬКО таблицу курса
        tbl_anc = p.xpath("ancestor::w:tbl[1]", namespaces=NS)
        if tbl_anc:
            if is_course_table(tbl_anc[0]):
                continue  # курс-таблица не трогаем, чтобы не ломать

        pPr = p.find("w:pPr", NS)
        if pPr is None:
            pPr = etree.Element(q("pPr"))
            p.insert(0, pPr)

        spacing = ensure(pPr, "spacing")
        spacing.set(q("before"), str(before))
        spacing.set(q("after"), str(after))
        spacing.set(q("line"), str(line))
        spacing.set(q("lineRule"), str(line_rule))

        cs = pPr.find("w:contextualSpacing", NS)
        if cs is not None:
            pPr.remove(cs)

        patched += 1

    files["word/document.xml"] = etree.tostring(
        root, xml_declaration=True, encoding="UTF-8", standalone="yes"
    )

    with zipfile.ZipFile(docx_path, "w", compression=zipfile.ZIP_DEFLATED) as zout:
        for n, data in files.items():
            zout.writestr(n, data)

    return patched


def _patch_docx_spacing_everywhere(
    docx_path: Path,
    before="0",
    after="0",
    line=240,          # 240 twips = 12pt “single” (примерно)
    line_rule="auto",  # auto vs exactly
):
    def q(tag): return f"{{{W_NS}}}{tag}"

    def ensure(parent, tag):
        el = parent.find(f"w:{tag}", NS)
        if el is None:
            el = etree.Element(q(tag))
            parent.append(el)
        return el

    with zipfile.ZipFile(docx_path, "r") as zin:
        files = {name: zin.read(name) for name in zin.namelist()}

    changed_files = 0
    changed_paras = 0

    targets = ["word/document.xml"]

    for name in targets:
        root = etree.fromstring(files[name])

        for p in root.xpath(".//w:p", namespaces=NS):
            pPr = p.find("w:pPr", NS)
            if pPr is None:
                pPr = etree.Element(q("pPr"))
                p.insert(0, pPr)

            ind = ensure(pPr, "ind")
            ind.set(q("left"), "0")
            ind.set(q("right"), "0")
            ind.set(q("firstLine"), "0")
            ind.set(q("hanging"), "0")

            spacing = ensure(pPr, "spacing")
            spacing.set(q("before"), str(before))
            spacing.set(q("after"), str(after))
            spacing.set(q("line"), str(line))
            spacing.set(q("lineRule"), str(line_rule))

            cs = pPr.find("w:contextualSpacing", NS)
            if cs is not None:
                pPr.remove(cs)

            changed_paras += 1

        files[name] = etree.tostring(
            root, xml_declaration=True, encoding="UTF-8", standalone="yes"
        )
        changed_files += 1

    with zipfile.ZipFile(docx_path, "w", compression=zipfile.ZIP_DEFLATED) as zout:
        for n, data in files.items():
            zout.writestr(n, data)

    logger.warning("DOCX: spacing patched files=%s paras=%s", changed_files, changed_paras)


def _patch_courseinfo_fonts_and_borders(
    docx_path: Path,
    markers=None,
    font_name="Trebuchet MS",
    font_size_pt=11,
    border_size=8,
    border_color="000000",
    bold_labels=True,
    row_height_cm=0.7,          # ✅ НОВОЕ
    v_align="center",           # ✅ НОВОЕ: center / top / bottom
):
    if markers is None:
        markers = COURSE_INFO_MARKERS

    half_points = str(int(font_size_pt * 2))

    def cm_to_twips(cm: float) -> str:
        tw = int(round((cm / 2.54) * 1440))
        return str(max(tw, 1))

    row_twips = cm_to_twips(row_height_cm)

    def q(tag):
        return f"{{{W_NS}}}{tag}"

    def ensure(parent, tag, insert0=False):
        el = parent.find(f"w:{tag}", NS)
        if el is None:
            el = etree.Element(q(tag))
            if insert0 and len(parent):
                parent.insert(0, el)
            else:
                parent.append(el)
        return el

    def set_run_style(run_el, make_bold=None):
        rpr = run_el.find("w:rPr", NS)
        if rpr is None:
            rpr = etree.Element(q("rPr"))
            run_el.insert(0, rpr)

        rfonts = rpr.find("w:rFonts", NS)
        if rfonts is None:
            rfonts = etree.Element(q("rFonts"))
            rpr.insert(0, rfonts)

        rfonts.set(q("ascii"), font_name)
        rfonts.set(q("hAnsi"), font_name)
        rfonts.set(q("eastAsia"), font_name)
        rfonts.set(q("cs"), font_name)

        sz = rpr.find("w:sz", NS)
        if sz is None:
            sz = etree.Element(q("sz"))
            rpr.append(sz)
        sz.set(q("val"), half_points)

        szcs = rpr.find("w:szCs", NS)
        if szcs is None:
            szcs = etree.Element(q("szCs"))
            rpr.append(szcs)
        szcs.set(q("val"), half_points)

        if make_bold is not None:
            if make_bold:
                ensure(rpr, "b")
            else:
                b = rpr.find("w:b", NS)
                if b is not None:
                    rpr.remove(b)

    def node_text(node):
        parts = node.xpath(".//w:t/text()", namespaces=NS)
        return " ".join(" ".join(parts).split()).strip()

    def is_course_table(tbl):
        found = set()
        for tr in tbl.xpath("./w:tr", namespaces=NS):
            tcs = tr.xpath("./w:tc", namespaces=NS)
            if not tcs:
                continue
            left = node_text(tcs[0])
            for m in markers:
                if left.startswith(m):
                    found.add(m)
        return len(found) >= 3

    def set_table_borders(tbl):
        tblPr = tbl.find("w:tblPr", NS)
        if tblPr is None:
            tblPr = etree.Element(q("tblPr"))
            tbl.insert(0, tblPr)

        borders = tblPr.find("w:tblBorders", NS)
        if borders is None:
            borders = etree.Element(q("tblBorders"))
            tblPr.append(borders)

        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            el = borders.find(f"w:{edge}", NS)
            if el is None:
                el = etree.Element(q(edge))
                borders.append(el)
            el.set(q("val"), "single")
            el.set(q("sz"), str(border_size))
            el.set(q("space"), "0")
            el.set(q("color"), border_color)

    def set_row_height(tr):
        trPr = tr.find("w:trPr", NS)
        if trPr is None:
            trPr = etree.Element(q("trPr"))
            tr.insert(0, trPr)

        trH = trPr.find("w:trHeight", NS)
        if trH is None:
            trH = etree.Element(q("trHeight"))
            trPr.append(trH)

        trH.set(q("val"), row_twips)
        trH.set(q("hRule"), "exact")   # ✅ именно EXACT

    def set_cell_valign(tc):
        tcPr = tc.find("w:tcPr", NS)
        if tcPr is None:
            tcPr = etree.Element(q("tcPr"))
            tc.insert(0, tcPr)

        vA = tcPr.find("w:vAlign", NS)
        if vA is None:
            vA = etree.Element(q("vAlign"))
            tcPr.append(vA)
        vA.set(q("val"), v_align)

        for p in tc.xpath(".//w:p", namespaces=NS):
            pPr = p.find("w:pPr", NS)
            if pPr is None:
                pPr = etree.Element(q("pPr"))
                p.insert(0, pPr)

            spacing = pPr.find("w:spacing", NS)
            if spacing is None:
                spacing = etree.Element(q("spacing"))
                pPr.append(spacing)
            spacing.set(q("before"), "0")
            spacing.set(q("after"), "0")

    with zipfile.ZipFile(docx_path, "r") as zin:
        files = {name: zin.read(name) for name in zin.namelist()}

    xml = files.get("word/document.xml")
    if not xml:
        return 0, 0

    root = etree.fromstring(xml)

    hits_tables = 0
    hits_paras = 0

    for tbl in root.xpath(".//w:tbl", namespaces=NS):
        if not is_course_table(tbl):
            continue

        hits_tables += 1
        set_table_borders(tbl)

        for tr in tbl.xpath("./w:tr", namespaces=NS):
            set_row_height(tr)
            for tc in tr.xpath("./w:tc", namespaces=NS):
                set_cell_valign(tc)

        for r in tbl.xpath(".//w:r", namespaces=NS):
            set_run_style(r, make_bold=None)

        if bold_labels:
            for tr in tbl.xpath("./w:tr", namespaces=NS):
                tcs = tr.xpath("./w:tc", namespaces=NS)
                if not tcs:
                    continue
                left_tc = tcs[0]
                left_txt = node_text(left_tc)
                if any(left_txt.startswith(m) for m in markers):
                    for r in left_tc.xpath(".//w:r", namespaces=NS):
                        set_run_style(r, make_bold=True)

        break

    for p in root.xpath(".//w:p", namespaces=NS):
        parts = p.xpath(".//w:t/text()", namespaces=NS)
        full = " ".join(" ".join(parts).split()).strip()
        if full and any(full.startswith(m) for m in markers):
            for r in p.xpath(".//w:r", namespaces=NS):
                set_run_style(r, make_bold=None)
            hits_paras += 1

    files["word/document.xml"] = etree.tostring(
        root, xml_declaration=True, encoding="UTF-8", standalone="yes"
    )

    with zipfile.ZipFile(docx_path, "w", compression=zipfile.ZIP_DEFLATED) as zout:
        for name, data in files.items():
            zout.writestr(name, data)

    return hits_tables, hits_paras


def _patch_docx_courseinfo_paragraphs(
    docx_path: Path,
    markers=None,
    font_name="Trebuchet MS",
    font_size_pt=11,
    bold_labels=True,   # ✅ ДОБАВИЛИ
):
    if markers is None:
        markers = globals().get("COURSE_INFO_MARKERS") or [
            "Nombre curso:",
            "Código curso:",
            "Año académico:",
            "Fechas impartición:",
            "Horario:",
            "Entidad:",
            "Tipo formación:",
        ]

    markers = [str(m).strip() for m in (markers or []) if str(m).strip()]
    half_points = str(int(font_size_pt * 2))

    def ensure_rpr(run_el):
        rpr = run_el.find("w:rPr", NS)
        if rpr is None:
            rpr = etree.Element(f"{{{W_NS}}}rPr")
            run_el.insert(0, rpr)
        return rpr

    def set_fonts(rpr):
        rfonts = rpr.find("w:rFonts", NS)
        if rfonts is None:
            rfonts = etree.Element(f"{{{W_NS}}}rFonts")
            rpr.insert(0, rfonts)
        rfonts.set(f"{{{W_NS}}}ascii", font_name)
        rfonts.set(f"{{{W_NS}}}hAnsi", font_name)
        rfonts.set(f"{{{W_NS}}}eastAsia", font_name)
        rfonts.set(f"{{{W_NS}}}cs", font_name)

        sz = rpr.find("w:sz", NS)
        if sz is None:
            sz = etree.Element(f"{{{W_NS}}}sz")
            rpr.append(sz)
        sz.set(f"{{{W_NS}}}val", half_points)

        szcs = rpr.find("w:szCs", NS)
        if szcs is None:
            szcs = etree.Element(f"{{{W_NS}}}szCs")
            rpr.append(szcs)
        szcs.set(f"{{{W_NS}}}val", half_points)

    def set_bold(rpr, on=True):
        b = rpr.find("w:b", NS)
        if b is None:
            b = etree.Element(f"{{{W_NS}}}b")
            rpr.append(b)
        b.set(f"{{{W_NS}}}val", "1" if on else "0")

    hits = 0

    with zipfile.ZipFile(docx_path, "r") as zin:
        files = {name: zin.read(name) for name in zin.namelist()}

    xml = files.get("word/document.xml")
    if not xml:
        logger.warning("DOCX XML: word/document.xml not found")
        return

    root = etree.fromstring(xml)

    for p in root.xpath(".//w:p", namespaces=NS):
        texts = [t.text for t in p.xpath(".//w:t", namespaces=NS) if t.text]
        full = " ".join(" ".join(texts).split()).strip()
        if not full:
            continue

        if any(full.startswith(m) for m in markers):
            runs = p.xpath(".//w:r", namespaces=NS)
            for r in runs:
                rpr = ensure_rpr(r)
                set_fonts(rpr)

            if bold_labels and runs:
                rpr0 = ensure_rpr(runs[0])
                set_bold(rpr0, True)

            hits += 1

    files["word/document.xml"] = etree.tostring(
        root, xml_declaration=True, encoding="UTF-8", standalone="yes"
    )

    with zipfile.ZipFile(docx_path, "w", compression=zipfile.ZIP_DEFLATED) as zout:
        for name, data in files.items():
            zout.writestr(name, data)

    logger.warning("DOCX XML: courseinfo paragraphs patched = %s", hits)


def _inject_headers_footers(docx_path: Path):
    LOGOS = getattr(settings, "MEATZE_DOCX_LOGOS", {})

    doc = Document(str(docx_path))
    section = doc.sections[0]

    section.left_margin  = Cm(1.0)
    section.right_margin = Cm(1.0)
    section.top_margin   = Cm(1.2)
    section.bottom_margin= Cm(1.2)

    usable_width = section.page_width - section.left_margin - section.right_margin

    # ================= HEADER =================
    header = section.header
    h_table = header.add_table(rows=1, cols=2, width=usable_width)
    col_left_width  = int(usable_width * 0.5)
    col_right_width = int(usable_width * 0.5)

    h_table.columns[0].width = col_left_width
    h_table.columns[1].width = col_right_width

    cell_left, cell_right = h_table.rows[0].cells

    lanbide = LOGOS.get("lanbide")
    if lanbide and Path(lanbide).exists():
        p_l = cell_left.paragraphs[0]
        p_l.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_l.paragraph_format.space_before = Pt(0)
        p_l.paragraph_format.space_after  = Pt(0)
        p_l.add_run().add_picture(lanbide, width=Cm(2.5))

    euskadi = LOGOS.get("euskadi")
    if euskadi and Path(euskadi).exists():
        p_r = cell_right.paragraphs[0]
        p_r.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_r.paragraph_format.space_before = Pt(0)
        p_r.paragraph_format.space_after  = Pt(0)
        p_r.add_run().add_picture(euskadi, width=Cm(9.5))

    # ================= FOOTER =================
    footer = section.footer
    for p in footer.paragraphs:
        p.text = ""

    section.footer_distance = Cm(1.2)

    f_table = footer.add_table(rows=1, cols=1, width=usable_width)
    f_table.autofit = False

    eu = LOGOS.get("eu")
    if eu and Path(eu).exists():
        p = f_table.rows[0].cells[0].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        p.add_run().add_picture(eu, width=Cm(3.2))

    doc.save(str(docx_path))


@csrf_exempt
@require_http_methods(["POST"])
@require_admin_token
def export_docx_graphic(request):
    try:
        payload = json.loads(request.body or "{}")
    except json.JSONDecodeError:
        return JsonResponse({"ok": False, "error": "json_invalid"}, status=400)

    codigo = (payload.get("codigo") or "").strip().upper()
    tipo   = (payload.get("tipo") or "curso").strip().lower()
    grupo  = (payload.get("grupo") or "").strip()
    vacaciones = payload.get("vacaciones") or []

    if not codigo:
        return JsonResponse({"ok": False, "error": "codigo_required"}, status=400)

    curso = get_object_or_404(Curso, codigo=codigo)

    html = (payload.get("html") or "").strip()

    if not html:
        header = build_horario_header_from_db(curso, tipo, grupo)
        mods   = build_modules_legend_from_db(curso)
        html = render_to_string("exports/calendario_docx.html", {
            "curso": curso,
            "header_main": header["main"],
            "header_exceptions": header["exceptions"],
            "modules": mods,
            "vacaciones": vacaciones,
        })

    filename = (payload.get("filename") or "calendario.docx").strip()
    if not filename.lower().endswith(".docx"):
        filename += ".docx"

    with tempfile.TemporaryDirectory() as tmpdir:
        tmpdir_path = Path(tmpdir)
        html_path = tmpdir_path / "calendario.html"

        html = (payload.get("html") or "").strip()
        html_from_payload = bool(html)

        if not html_from_payload:
            try:
                mods = build_modules_legend_from_db(curso)
                legend_html = render_to_string(
                    "exports/_modules_legend.html",
                    {"mods": mods, "curso": curso}
                )

                if "<!--MZ_MODULES_LEGEND-->" in html:
                    html = html.replace("<!--MZ_MODULES_LEGEND-->", legend_html)
                    logger.warning("EXPORT: modules legend injected via marker")
                else:
                    if "</body>" in html:
                        html = html.replace("</body>", legend_html + "\n</body>", 1)
                    elif "</html>" in html:
                        html = html.replace("</html>", legend_html + "\n</html>", 1)
                    else:
                        html += "\n" + legend_html
            except Exception:
                logger.exception("Failed to inject modules legend into HTML")

        logger.warning("EXPORT html_len=%s head=%r", len(html or ""), (html or "")[:400])
        html_path.write_text(html, encoding="utf-8")

        if not LIBREOFFICE_BIN or not Path(LIBREOFFICE_BIN).exists():
            return JsonResponse(
                {"ok": False, "error": "soffice_not_found", "bin": LIBREOFFICE_BIN},
                status=503
            )

        cmd = [
            LIBREOFFICE_BIN,
            "--headless", "--nologo",
            "--convert-to", "docx:MS Word 2007 XML",
            "--outdir", str(tmpdir_path),
            str(html_path),
        ]

        env = os.environ.copy()
        env["PATH"] = env.get("PATH") or "/usr/local/sbin:/usr/local/bin:/usr/sbin:/usr/bin:/sbin:/bin"

        proc = subprocess.run(
            cmd,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            cwd=tmpdir,
            timeout=120,
            env=env,
        )
        logger.info("LO stdout: %s", proc.stdout.decode(errors="ignore"))
        logger.info("LO stderr: %s", proc.stderr.decode(errors="ignore"))

        if proc.returncode != 0:
            return JsonResponse(
                {
                    "ok": False,
                    "error": "soffice_failed",
                    "stdout": proc.stdout.decode(errors="ignore"),
                    "stderr": proc.stderr.decode(errors="ignore"),
                },
                status=500,
            )

        docx_files = list(tmpdir_path.glob("*.docx"))
        if not docx_files:
            return JsonResponse({"ok": False, "error": "docx_not_generated"}, status=500)

        docx_path = docx_files[0]

        try:
            _inject_headers_footers(docx_path)
        except Exception:
            logger.exception("Failed to inject headers/footers")

        try:
            _patch_docx_courseinfo_paragraphs(docx_path, font_size_pt=11)
        except Exception:
            logger.exception("DOCX patch failed, continuing")

        doc = Document(str(docx_path))

        try:
            _add_legend_nonlective(doc, vacaciones)
        except Exception:
            logger.exception("Failed to add nonlective legend")

        doc.save(str(docx_path))
        try:
            _patch_courseinfo_fonts_and_borders(
                docx_path,
                font_size_pt=12,
                row_height_cm=0.7,
                v_align="center",
            )
        except Exception:
            logger.exception("DOCX border patch failed")
        # XML-бодер патч (если он “про таблицу курса”)
        try:
            # 1) сжать всё после "Leyenda de módulos" (включая список модулей)
            _patch_docx_spacing_between_anchors(
                docx_path,
                start_anchor="Leyenda de módulos",
                end_anchor=None,
                before="0", after="0",
                line=240, line_rule="auto",
            )
        except Exception:
            logger.exception("DOCX legend spacing patch failed")
        except Exception as e:
            logger.exception("Failed to add nonlective legend: %s", e)

        data = docx_path.read_bytes()

    resp = HttpResponse(
        data,
        content_type=(
            "application/vnd.openxmlformats-officedocument."
            "wordprocessingml.document"
        ),
    )
    resp["Content-Disposition"] = f"attachment; filename=\"{filename}\""
    return resp