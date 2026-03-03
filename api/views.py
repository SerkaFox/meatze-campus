from django.shortcuts import render, redirect
from django.utils.http import url_has_allowed_host_and_scheme
from rest_framework.response import Response
from django.conf import settings
from django.http import JsonResponse, HttpResponse
from django.views.decorators.csrf import csrf_exempt
from rest_framework.decorators import api_view, permission_classes, authentication_classes
from django.views.decorators.http import require_http_methods
from rest_framework.permissions import AllowAny
from django.contrib.auth import get_user_model, authenticate, login
import json, io, re, os, requests, subprocess, tempfile, logging, random
from rest_framework import status
from pathlib import Path
from functools import wraps
from .models import Curso, Enrol, UserProfile, Horario, PendingRole, LoginPIN, MZSetting
from django.db import transaction
from django.core.cache import cache
from django.core.mail import send_mail
from datetime import date, time, timedelta, datetime
from django.db import transaction
from io import BytesIO
from docx import Document 
from docx.shared import Mm, Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from django.db.models import Q, Count
from django.shortcuts import get_object_or_404
from .decorators import require_admin_token 
from panel.models import (CursoFile, CursoPhysicalConfig, MaterialDownload, MaterialReceipt)
from panel.views import PHYSICAL_ITEMS  
from django.template.loader import render_to_string
from api.horario_legend import (
    build_horario_header_from_db,
    build_modules_legend_from_db,
    _add_legend_nonlective,
    _add_modules_legend_blocks,
)
logger = logging.getLogger(__name__)

LIBREOFFICE_BIN = getattr(settings, "LIBREOFFICE_PATH", "/usr/bin/soffice")

MEATZE_DOCX_LOGOS = getattr(settings, "MEATZE_DOCX_LOGOS", {})


User = get_user_model()

@api_view(["GET"])
@permission_classes([AllowAny])
@authentication_classes([])
def public_cursos_for_temp(request):
    qs = Curso.objects.all().order_by("codigo")

    items = []
    for c in qs:
        items.append({
            "codigo": c.codigo,
            "titulo": c.titulo,
        })

    return Response({"items": items})


@csrf_exempt
@api_view(['POST'])
@permission_classes([AllowAny])
@authentication_classes([])
def login_password(request):
    data = {}
    try:
        if isinstance(request.data, dict):
            data.update(request.data)
    except Exception as e:
        logger.warning("login_password: error reading request.data: %s", e)
    if not data:
        try:
            raw = request.body
            if isinstance(raw, (bytes, bytearray)):
                raw = raw.decode('utf-8', errors='ignore')
            raw = (raw or '').strip()
            if raw:
                data.update(json.loads(raw))
        except Exception as e:
            logger.warning(
                "login_password: cannot parse raw body as JSON: %s; body=%r",
                e, request.body[:200]
            )

    email = (data.get('email') or '').strip().lower()
    password = (data.get('password') or '').strip()

    if not email or not password:
        logger.warning(
            "login_password: missing email/password. data=%r body=%r",
            data, request.body[:200]
        )
        return Response(
            {'message': 'Falta e-mail o contraseña'},
            status=status.HTTP_400_BAD_REQUEST
        )

    # 2) логиним — у тебя username == email, судя по коду создания юзеров
    user = authenticate(request, username=email, password=password)

    if user is None:
        return Response(
            {'message': 'Credenciales inválidas.'},
            status=status.HTTP_400_BAD_REQUEST
        )

    # 3) создаём сессию
    login(request, user)

    # 4) собираем "me" для фронта
    me = {
        'id': user.id,
        'email': getattr(user, 'email', '') or '',
        'first_name': getattr(user, 'first_name', ''),
        'last_name': getattr(user, 'last_name', ''),
        'is_teacher': getattr(user, 'is_staff', False),  # у тебя docente = staff
        'has_password': True,
    }

    return Response({'me': me}, status=status.HTTP_200_OK)


def require_admin_token(view_func):
    @wraps(view_func)
    def _wrapped(request, *args, **kwargs):
        token = request.GET.get("adm") or request.headers.get("X-MZ-Admin")
        expected = getattr(settings, "MEATZE_ADMIN_PASS", "MeatzeIT")

        # 👇 временный лог
        print("ADM DEBUG:", repr(token), "expected:", repr(expected))

        if token != expected:
            return JsonResponse({"ok": False, "error": "forbidden"}, status=403)
        return view_func(request, *args, **kwargs)
    return _wrapped



@csrf_exempt
@require_admin_token
def admin_ping(request):
    # Просто проверка токена
    return JsonResponse({"ok": True})


@require_admin_token
@require_http_methods(["GET"])
def admin_cursos_list(request):
    qs = Curso.objects.all().order_by("codigo")
    items = []

    for c in qs:
        # в модели modules — LIST
        modules_str = json.dumps(c.modules or [], ensure_ascii=False)

        items.append({
            "id": c.id,
            "codigo": c.codigo,
            "titulo": c.titulo,
            "modules": modules_str,
            "horas": c.horas_total or 0,
            "tipo_formacion": c.tipo_formacion,
        })

    return JsonResponse({"items": items})

def _admin_can_access_course(request, curso: Curso) -> bool:
    """
    ВАЖНО: тут решаешь политику.
    Минимальный вариант: любой валидный X-MZ-Admin видит всё.
    Лучше: проверить, что admin связан с курсом/организацией.
    """
    # Вариант A (просто): True
    return True

    # Вариант B (пример): токен связан с user (если require_admin_token кладет request.admin_user)
    # u = getattr(request, "admin_user", None)
    # if not u:
    #     return False
    # return Enrol.objects.filter(user=u, codigo=curso.codigo, role__in=["teacher","admin"]).exists()


@require_admin_token
@require_http_methods(["GET"])
def admin_material_status(request):
    """
    Возвращает тот же формат что alumnos_status, но auth по X-MZ-Admin
    GET /meatze/v5/admin/lanbide/material_status?codigo=IFCT0309
    """
    codigo = (request.GET.get("codigo") or "").strip().upper()
    if not codigo:
        return JsonResponse({"ok": False, "error": "codigo required"}, status=400)

    curso = get_object_or_404(Curso, codigo=codigo)

    if not _admin_can_access_course(request, curso):
        return JsonResponse({"ok": False, "error": "forbidden"}, status=403)

    # --- физические предметы курса ---
    cfg = CursoPhysicalConfig.objects.filter(curso=curso).only("enabled_keys").first()
    physical_keys = (cfg.enabled_keys if cfg else []) or []
    total_phys = len(physical_keys)

    # --- файлы, видимые ученикам ---
    visible_files_qs = (
        CursoFile.objects.filter(curso=curso)
        .filter(
            Q(tipo=CursoFile.TIPO_ALUMNOS) |
            (Q(share_alumnos=True) & Q(tipo__in=[CursoFile.TIPO_DOCENTES, CursoFile.TIPO_PRIVADO]))
        )
        .only("id", "title", "file")
    )

    visible_ids = list(visible_files_qs.values_list("id", flat=True))
    total_visible = len(visible_ids)

    # если файлов нет и физики нет — быстрый ответ
    if total_visible == 0 and total_phys == 0:
        return JsonResponse({
            "ok": True,
            "codigo": curso.codigo,
            "total_files": 0,
            "total_phys": 0,
            "items": {}
        })

    # --- агрегаты ---
    dl_counts = dict(
        MaterialDownload.objects.filter(file_id__in=visible_ids)
        .values("alumno_id").annotate(c=Count("id"))
        .values_list("alumno_id", "c")
    )

    rc_counts = dict(
        MaterialReceipt.objects.filter(curso=curso, item_key__in=physical_keys)
        .values("alumno_id").annotate(c=Count("id"))
        .values_list("alumno_id", "c")
    )

    # --- мапа file_id -> title/filename ---
    file_map = {}
    for f in visible_files_qs:
        # f.file может быть FileField; у тебя раньше было f.filename, но не всегда есть
        filename = ""
        try:
            filename = getattr(f.file, "name", "") or ""
            filename = filename.rsplit("/", 1)[-1]
        except Exception:
            filename = ""
        file_map[f.id] = (f.title or filename or f"#{f.id}")

    # --- детали: downloads ---
    dl_detail = {}
    dl_rows = (
        MaterialDownload.objects
        .filter(file_id__in=visible_ids)
        .values_list("alumno_id", "file_id")
    )
    for aid, fid in dl_rows:
        dl_detail.setdefault(aid, []).append(file_map.get(fid, f"#{fid}"))

    # --- детали: receipts ---
    label_by_key = {it["key"]: it["label"] for it in PHYSICAL_ITEMS}
    rc_detail = {}
    rc_rows = (
        MaterialReceipt.objects
        .filter(curso=curso, item_key__in=physical_keys)
        .values_list("alumno_id", "item_label", "item_key")
    )
    for aid, item_label, item_key in rc_rows:
        label = (item_label or "").strip() or label_by_key.get(item_key, item_key)
        rc_detail.setdefault(aid, []).append(label)

    # --- все alumno_id ---
    all_ids = set(dl_counts) | set(rc_counts) | set(dl_detail) | set(rc_detail)

    # (опционально) чистим дубликаты в деталях, чтобы tooltip был аккуратней
    def _uniq(seq):
        seen = set()
        out = []
        for x in seq:
            if x in seen:
                continue
            seen.add(x)
            out.append(x)
        return out

    items = {}
    for aid in all_ids:
        items[str(aid)] = {
            "d": int(dl_counts.get(aid, 0)),
            "r": int(rc_counts.get(aid, 0)),
            "dl": _uniq(dl_detail.get(aid, [])),
            "rc": _uniq(rc_detail.get(aid, [])),
        }

    return JsonResponse({
        "ok": True,
        "codigo": curso.codigo,
        "total_files": total_visible,
        "total_phys": total_phys,
        "items": items,
    })




@api_view(["GET"])
def curso_detail(request):
    """
    /meatze/v5/curso?codigo=...
    Один курс по коду.
    """
    codigo = request.query_params.get("codigo")
    if not codigo:
        return Response({"error": "codigo requerido"}, status=400)

    try:
        c = Curso.objects.get(codigo=codigo)
    except Curso.DoesNotExist:
        return Response({"error": "not_found"}, status=404)

    data = {
        "codigo": c.codigo,
        "titulo": c.titulo,
        "descripcion": c.descripcion,
        "fecha_inicio": c.fecha_inicio,
        "fecha_fin": c.fecha_fin,
    }
    return Response({"curso": data})


def acceder(request):
    tab = (request.GET.get("tab") or "login").strip().lower()

    # ✅ ВАЖНО: если пользователь залогинен и просит профиль — НЕ редиректим
    if request.user.is_authenticated and tab == "profile":
        return render(request, "acceder.html", {})

    # (опционально) если залогинен и просто открыл /acceder/ без tab
    # отправляем в панель
    if request.user.is_authenticated and tab in ("", "login"):
        return redirect("/alumno/")

    # не залогинен → показываем страницу входа/регистрации
    return render(request, "acceder.html", {})

# ─────────────────────────────
# ADMIN · DOCENTES
# ─────────────────────────────

def _split_last_name(last_name: str):
    last_name = (last_name or "").strip()
    if not last_name:
        return "", ""
    parts = last_name.split(" ", 1)
    if len(parts) == 1:
        return parts[0], ""
    return parts[0], parts[1]


@require_admin_token
@require_http_methods(["GET"])
def admin_teachers_list(request):
    qs = (
        User.objects.filter(is_staff=True)
        .select_related("profile")
        .order_by("first_name", "last_name", "email")
    )

    users = list(qs)
    user_ids = [u.id for u in users]

    # все назначения teacher по этим пользователям
    enrols = (
        Enrol.objects
        .filter(user_id__in=user_ids, role="teacher")
        .values_list("user_id", "codigo")
    )

    by_user = {}
    codigos = set()
    for uid, codigo in enrols:
        codigo = (codigo or "").strip().upper()
        if not codigo:
            continue
        by_user.setdefault(uid, []).append(codigo)
        codigos.add(codigo)

    # подтянем названия курсов
    titulo_by_codigo = dict(
        Curso.objects.filter(codigo__in=list(codigos))
        .values_list("codigo", "titulo")
    )

    items = []
    for u in users:
        profile = getattr(u, "profile", None)
        wa = (profile.wa if profile else "") or ""
        first_name = (profile.first_name if profile and profile.first_name else u.first_name or "").strip()
        last_name1 = (profile.last_name1 if profile else "") or ""
        last_name2 = (profile.last_name2 if profile else "") or ""
        display_name = (
            profile.display_name
            if profile and profile.display_name
            else (u.get_full_name() or u.username or u.email)
        )

        cods = by_user.get(u.id, []) or []
        cursos = [{"codigo": c, "titulo": (titulo_by_codigo.get(c) or "")} for c in sorted(set(cods))]

        items.append({
            "id": u.id,
            "email": u.email or u.username,
            "wa": wa,
            "first_name": first_name,
            "last_name1": last_name1,
            "last_name2": last_name2,
            "display_name": display_name,
            "cursos": cursos,          # ✅ НОВОЕ
        })

    return JsonResponse({"items": items})
    
    
from django.db import transaction, IntegrityError

@require_admin_token
@csrf_exempt
@require_http_methods(["GET", "POST"])
def admin_teachers_upsert(request):
    if request.method == "GET":
        return admin_teachers_list(request)

    try:
        data = json.loads(request.body or "{}")
    except json.JSONDecodeError:
        return JsonResponse({"ok": False, "message": "json_invalid"}, status=400)

    user_id = data.get("id")
    email = (data.get("email") or "").strip().lower()
    if not email:
        return JsonResponse({"ok": False, "message": "email_required"}, status=400)

    first_name = (data.get("first_name") or "").strip()
    last1 = (data.get("last_name1") or "").strip()
    last2 = (data.get("last_name2") or "").strip()
    bio  = (data.get("bio") or "").strip()
    wa   = (data.get("wa") or "").strip()
    full_last = " ".join(p for p in [last1, last2] if p).strip()

    with transaction.atomic():
        user = None

        # =========================
        # A) UPDATE по ID
        # =========================
        if user_id:
            user = User.objects.filter(pk=int(user_id)).first()
            if not user:
                return JsonResponse({"ok": False, "message": "not_found"}, status=404)

            current_email = ((user.email or user.username or "")).strip().lower()

            # ❌ запрещаем менять email (и НЕ трогаем username)
            if email and email != current_email:
                return JsonResponse({
                    "ok": False,
                    "message": "email_change_forbidden",
                    "detail": (
                        "No se puede cambiar el e-mail de un docente existente. "
                        "Elimínalo y vuelve a añadirlo con el nuevo e-mail."
                    )
                }, status=409)

        # =========================
        # B) UPSERT по email (CREATE/PROMOTE)
        # =========================
        if not user:
            # 1) сначала ищем по username=email (это твой уникальный ключ)
            user = User.objects.filter(username__iexact=email).first()

            # 2) если не нашли — ищем по email
            if not user:
                user = User.objects.filter(email__iexact=email).first()

            # 3) если нашли по email, но username другой — НЕ пытаемся менять username
            #    это и вызывает UniqueViolation. Лучше вернуть 409 с инструкцией.
            if user and (user.username or "").strip().lower() != email:
                return JsonResponse({
                    "ok": False,
                    "message": "username_mismatch",
                    "detail": (
                        "Ese e-mail ya existe en el sistema con otro identificador. "
                        "Para cambiarlo, elimina el usuario anterior (purge) y vuelve a crearlo."
                    )
                }, status=409)

            # 4) если всё ещё нет — создаём (с защитой от гонки)
            if not user:
                try:
                    user = User.objects.create(username=email, email=email)
                except IntegrityError:
                    # на случай параллельного запроса — ещё раз читаем
                    user = User.objects.filter(username__iexact=email).first()
                    if not user:
                        raise

            # 5) если user найден по username=email, но email у него был "deleted+...@invalid"
            #    можно спокойно вернуть email обратно (email не уникальный обычно)
            if (user.email or "").strip().lower() != email:
                user.email = email

        # --- дальше обновляем поля преподавателя ---
        user.first_name = first_name
        user.last_name = full_last
        user.is_staff = True
        user.save()

        profile, _ = UserProfile.objects.get_or_create(user=user)
        profile.first_name = first_name
        profile.last_name1 = last1
        profile.last_name2 = last2
        profile.bio = bio
        profile.wa = wa
        if (data.get("display_name") or "").strip():
            profile.display_name = (data.get("display_name") or "").strip()
        profile.is_teacher = True if hasattr(profile, "is_teacher") else getattr(profile, "is_teacher", False)
        profile.save()

    return JsonResponse({"ok": True, "id": user.id})
    
    
@require_admin_token
@csrf_exempt
@require_http_methods(["POST"])
def admin_teacher_update(request, user_id: int):
    """
    POST /meatze/v5/admin/teachers/<id>
    Обновление данных преподавателя.
    """
    try:
        teacher = User.objects.get(pk=user_id)
    except User.DoesNotExist:
        return JsonResponse({"ok": False, "message": "not_found"}, status=404)

    try:
        data = json.loads(request.body or "{}")
    except json.JSONDecodeError:
        return JsonResponse({"ok": False, "message": "json_invalid"}, status=400)

    email = (data.get("email") or "").strip().lower()
    wa  = (data.get("wa") or "").strip()
    first_name = (data.get("first_name") or "").strip()
    last1 = (data.get("last_name1") or "").strip()
    last2 = (data.get("last_name2") or "").strip()
    bio = (data.get("bio") or "").strip()
    full_last = " ".join(p for p in [last1, last2] if p).strip()

    # --- проверяем конфликт по email, если его изменили ---
    current_email = (teacher.email or "").lower()
    if email and email != current_email:
        conflict = User.objects.filter(email=email).exclude(pk=teacher.pk).first()
        if conflict:
            # другой уже препод
            if conflict.is_staff:
                return JsonResponse(
                    {"ok": False, "message": "email_exists"},
                    status=409
                )
            # почта занята другим пользователем (alumno и т.п.)
            return JsonResponse(
                {"ok": False, "message": "email_in_use"},
                status=409
            )

        # если конфликта нет — обновляем email
        teacher.email = email
        if not teacher.username:
            teacher.username = email

    teacher.first_name = first_name
    teacher.last_name = full_last
    teacher.is_staff = True   # считаем staff = docente
    teacher.save()

    profile, _ = UserProfile.objects.get_or_create(user=teacher)
    profile.first_name = first_name
    profile.wa = wa
    profile.last_name1 = last1
    profile.last_name2 = last2
    profile.display_name = (
        data.get("display_name")
        or f"{first_name} {last1}".strip()
        or email
        or teacher.username
    )
    profile.bio = bio
    profile.save()

    return JsonResponse({"ok": True})


@require_admin_token
@csrf_exempt
@require_http_methods(["POST"])
def admin_teacher_delete(request, user_id: int):
    """
    POST /meatze/v5/admin/teachers/<id>/delete
    Убираем статус staff и чистим Enrol с ролью teacher.
    """
    try:
        user = User.objects.get(pk=user_id)
    except User.DoesNotExist:
        return JsonResponse({"message": "not_found"}, status=404)

    # перестаёт быть docente
    user.is_staff = False
    user.save(update_fields=["is_staff"])

    # на всякий случай чистим привязки как преподавателя
    Enrol.objects.filter(user=user, role="teacher").delete()

    return JsonResponse({"ok": True})


@csrf_exempt
@require_admin_token
def admin_cursos_delete(request, curso_id):
    """
    POST /meatze/v5/admin/cursos/<id>/delete
    """
    if request.method != "POST":
        return JsonResponse({"error": "method_not_allowed"}, status=405)

    try:
        curso = Curso.objects.get(id=curso_id)
    except Curso.DoesNotExist:
        return JsonResponse({"error": "not_found"}, status=404)

    curso.delete()
    return JsonResponse({"ok": True})


@require_admin_token
@csrf_exempt
@require_http_methods(["POST"])
def admin_cursos_upsert(request):
    try:
        data = json.loads(request.body or "{}")
    except json.JSONDecodeError:
        return JsonResponse({"message": "json_invalid"}, status=400)

    codigo = (data.get("codigo") or "").strip().upper()
    titulo = (data.get("titulo") or "").strip()
    modules = data.get("modules") or []
    tipo_formacion = (data.get("tipo_formacion") or "").strip().lower()

    # ✅ NEW: откуда клонировать материалы
    clone_from = (data.get("clone_from") or "").strip().upper()

    if not codigo or not titulo:
        return JsonResponse({"message": "codigo_titulo_required"}, status=400)

    norm_modules = []
    total_horas = 0
    for m in modules:
        if not isinstance(m, dict):
            continue
        name = (m.get("name") or "").strip()
        if not name:
            continue
        hours = int(m.get("hours") or 0)
        if hours < 0:
            hours = 0
        total_horas += hours
        norm_modules.append({"name": name, "hours": hours})

    curso_id = data.get("id")

    # определяем: это создание нового курса или апдейт
    creating = False

    with transaction.atomic():
        if curso_id:
            # редактирование по id
            try:
                curso = Curso.objects.get(id=curso_id)
            except Curso.DoesNotExist:
                # если id битый — fallback
                curso = Curso.objects.filter(codigo=codigo).first()
                if not curso:
                    curso = Curso(codigo=codigo)
                    creating = True
        else:
            # если нет id — upsert по коду
            curso = Curso.objects.filter(codigo=codigo).first()
            if not curso:
                curso = Curso(codigo=codigo)
                creating = True

        curso.codigo = codigo
        curso.titulo = titulo
        curso.modules = norm_modules
        curso.horas_total = total_horas
        curso.tipo_formacion = tipo_formacion
        curso.save()

        # ✅ клонируем материалы ТОЛЬКО при создании нового курса
        cloned_files = 0
        if creating and clone_from and clone_from != codigo:
            src = Curso.objects.filter(codigo=clone_from).first()
            if src:
                # переносим записи CursoFile без копирования файла на диск
                src_qs = (CursoFile.objects
                          .filter(curso=src)
                          .only("tipo", "module_key", "title", "file", "size", "ext", "share_alumnos", "uploaded_by", "folder_path"))

                batch = []
                for f in src_qs:
                    nf = CursoFile(
                        curso=curso,
                        uploaded_by=f.uploaded_by,
                        tipo=f.tipo,
                        module_key=f.module_key,
                        title=f.title,
                        size=f.size,
                        ext=f.ext,
                        share_alumnos=f.share_alumnos,
                        folder_path=f.folder_path,
                    )
                    nf.file.name = f.file.name  # ✅ та же ссылка
                    batch.append(nf)

                if batch:
                    CursoFile.objects.bulk_create(batch)
                    cloned_files = len(batch)

    return JsonResponse({
        "ok": True,
        "id": curso.id,
        "horas": total_horas,
        "tipo_formacion": curso.tipo_formacion,
        "cloned_files": cloned_files,   # ✅ чтобы фронт мог показать “скопировано N”
    })



@require_admin_token
@csrf_exempt
@require_http_methods(["POST"])
def admin_curso_delete(request, curso_id: int):
    """
    POST /meatze/v5/admin/cursos/<id>/delete
    """
    try:
        curso = Curso.objects.get(pk=curso_id)
    except Curso.DoesNotExist:
        return JsonResponse({"message": "not_found"}, status=404)

    codigo = curso.codigo
    curso.delete()
    Enrol.objects.filter(codigo=codigo).delete()

    return JsonResponse({"ok": True})


# ─────────────────────────────
# ADMIN · ENROLMENTS (teachers per curso)
# ─────────────────────────────

@require_admin_token
@require_http_methods(["GET"])
def admin_enrolments_list(request):
    """
    GET /meatze/v5/admin/enrolments?codigo=IFCT0209&role=teacher
    Для вкладки "Asignar docentes".
    """
    codigo = (request.GET.get("codigo") or "").strip()
    role = (request.GET.get("role") or "").strip()

    qs = Enrol.objects.select_related("user").all()
    if codigo:
        qs = qs.filter(codigo=codigo)
    if role:
        qs = qs.filter(role=role)

    items = []
    for e in qs:
        u = e.user
        items.append({
            "id": e.id,
            "user_id": u.id,
            "email": u.email or u.username,
            "display_name": u.get_full_name() or (u.email or u.username),
            "role": e.role,
            "codigo": e.codigo,
        })

    return JsonResponse({"items": items})


@require_admin_token
@csrf_exempt
@require_http_methods(["POST"])
def admin_cursos_assign(request):
    """
    POST /meatze/v5/admin/cursos/assign
    тело: {curso_codigo: "...", teachers: ["1","2",...]}
    Обновляет список Enrol с ролью teacher для данного курса.
    """
    try:
        data = json.loads(request.body or "{}")
    except json.JSONDecodeError:
        return JsonResponse({"message": "json_invalid"}, status=400)

    codigo = (data.get("curso_codigo") or "").strip()
    raw_ids = data.get("teachers") or []
    if not codigo:
        return JsonResponse({"message": "codigo_required"}, status=400)

    try:
        Curso.objects.get(codigo=codigo)
    except Curso.DoesNotExist:
        return JsonResponse({"message": "curso_not_found"}, status=404)

    teacher_ids = []
    for v in raw_ids:
        try:
            teacher_ids.append(int(v))
        except (TypeError, ValueError):
            continue

    teacher_ids = list(set(teacher_ids))

    with transaction.atomic():
        # удаляем все старые привязки teacher для этого курса
        Enrol.objects.filter(codigo=codigo, role="teacher").exclude(
            user_id__in=teacher_ids
        ).delete()

        # существующие
        existing = set(
            Enrol.objects.filter(codigo=codigo, role="teacher")
            .values_list("user_id", flat=True)
        )

        for uid in teacher_ids:
            if uid in existing:
                continue
            try:
                u = User.objects.get(pk=uid)
            except User.DoesNotExist:
                continue
            Enrol.objects.create(user=u, codigo=codigo, role="teacher")

    return JsonResponse({"ok": True, "teachers": teacher_ids})


# ─────────────────────────────
# STUBS для частей, которые пока не реализованы (чтобы не было 404)
# ─────────────────────────────
@csrf_exempt
@require_admin_token
@require_http_methods(["GET", "POST"])
def admin_fixed_nonlective(request):
    """
    GET  /meatze/v5/admin/fixed-nonlective?adm=...
         → {"years": {"2025": ["01-13","08-15", ...], ...}}

    POST /meatze/v5/admin/fixed-nonlective?adm=...
         тело: {"years": {"2025": "13/01, 15/01-20/01", "2026": "10/01, 20/06"}}
    """
    CACHE_KEY = "mz_fixed_nonlective"

    if request.method == "GET":
        years = cache.get(CACHE_KEY)
        if years is None:
            cfg = MZSetting.objects.filter(key="fixed_nonlective").first()
            years = (cfg.value if cfg else {}) or {}
            cache.set(CACHE_KEY, years, None)

        norm = {}
        for y, v in (years or {}).items():
            if isinstance(v, list):
                norm[str(y)] = [str(x).strip() for x in v if str(x).strip()]
            elif isinstance(v, str):
                norm[str(y)] = [s.strip() for s in v.split(",") if s.strip()]
            else:
                norm[str(y)] = []
        return JsonResponse({"years": norm})

    # ----- POST -----
    # ✅ 1) parse JSON body
    try:
        raw = (request.body or b"").decode("utf-8").strip()
        payload = json.loads(raw) if raw else {}
    except Exception:
        return JsonResponse({"error": "invalid_json"}, status=400)

    years_raw = payload.get("years") or {}
    if not isinstance(years_raw, dict):
        return JsonResponse({"error": "years_invalid"}, status=400)

    def expand_token_to_dates(y: int, token: str) -> list[str]:
        token = token.strip()
        if not token:
            return []

        # 01/08-31/08
        m = re.match(r"^(\d{1,2})/(\d{1,2})\s*-\s*(\d{1,2})/(\d{1,2})$", token)
        if m:
            d1, mo1, d2, mo2 = map(int, m.groups())
            try:
                start = date(y, mo1, d1)
                end = date(y, mo2, d2)
            except ValueError:
                return []
            if end < start:
                return []
            out, cur = [], start
            while cur <= end:
                out.append(cur.strftime("%m-%d"))
                cur += timedelta(days=1)
            return out

        # 10-20/09
        m = re.match(r"^(\d{1,2})\s*-\s*(\d{1,2})/(\d{1,2})$", token)
        if m:
            d1, d2, mo = map(int, m.groups())
            try:
                start = date(y, mo, d1)
                end = date(y, mo, d2)
            except ValueError:
                return []
            if end < start:
                return []
            out, cur = [], start
            while cur <= end:
                out.append(cur.strftime("%m-%d"))
                cur += timedelta(days=1)
            return out

        # 13/01
        m = re.match(r"^(\d{1,2})/(\d{1,2})$", token)
        if m:
            d, mo = map(int, m.groups())
            try:
                dt = date(y, mo, d)
            except ValueError:
                return []
            return [dt.strftime("%m-%d")]

        return []

    new_years: dict[str, list[str]] = {}

    for y_str, spec in years_raw.items():
        try:
            y = int(str(y_str))
        except ValueError:
            continue

        if isinstance(spec, list):
            tokens = [str(x) for x in spec]
        else:
            tokens = [t.strip() for t in str(spec).split(",") if t.strip()]

        mmdd_set = set()
        for t in tokens:
            for mmdd in expand_token_to_dates(y, t):
                mmdd_set.add(mmdd)

        new_years[str(y)] = sorted(mmdd_set)

    MZSetting.objects.update_or_create(
        key="fixed_nonlective",
        defaults={"value": new_years},
    )
    cache.set(CACHE_KEY, new_years, None)

    return JsonResponse({"ok": True, "years": new_years})

@require_admin_token
@require_http_methods(["GET"])
def admin_holidays(request, year: int):
    """
    GET /meatze/v5/admin/holidays/<year>?adm=...
    Возвращает кэшированный список праздников ES/ES-PV для фронта.

    Формат ответа:
    {
      "year": 2025,
      "items": [
        {"date": "2025-01-01", "name": "Año Nuevo"},
        ...
      ]
    }
    """

    year = int(year)
    CACHE_KEY = f"mz_holidays_{year}"

    # 1) пробуем взять из кеша
    cached = cache.get(CACHE_KEY)
    if cached is not None:
        return JsonResponse({"year": year, "items": cached})

    # 2) если нет в кеше — тянем из Nager.Date
    items = []
    try:
        r = requests.get(
            f"https://date.nager.at/api/v3/PublicHolidays/{year}/ES",
            timeout=10,
        )
        r.raise_for_status()
        raw = r.json()

        for h in raw:
            counties = h.get("counties") or []
            # если список провинций есть, берём только те, где есть ES-PV
            if counties and "ES-PV" not in counties:
                continue

            items.append({
                "date": h.get("date"),                          # YYYY-MM-DD
                "name": h.get("localName") or h.get("name", ""),  # локальное имя
            })

    except Exception as e:
        # на всякий случай не роняем фронт — просто пустой список
        print("HOLIDAYS ERROR:", e)
        items = []

    # 3) кладём в кеш, например, на сутки/месяц (можно и без TTL — на весь год)
    cache.set(CACHE_KEY, items, 60 * 60 * 24 * 30)  # 30 дней

    return JsonResponse({"year": year, "items": items})



@require_admin_token
@require_http_methods(["GET"])
def news_subscribers(request):
    """GET /meatze/v5/news/subscribers — заглушка."""
    return JsonResponse({"items": []})


@require_admin_token
@require_http_methods(["GET"])
def news_wa_inbox(request):
    """GET /meatze/v5/news/wa-inbox — заглушка."""
    return JsonResponse({"items": []})
    
@csrf_exempt
@require_http_methods(["GET", "POST"])
def curso_horario(request, codigo):
    """
    GET  /meatze/v5/curso/<codigo>/horario?tipo=curso|practica&grupo=...
    POST /meatze/v5/curso/<codigo>/horario?adm=...
         тело: {items:[...], tipo?, grupo?}
    """
    try:
        curso = Curso.objects.get(codigo=codigo)
    except Curso.DoesNotExist:
        return JsonResponse({"error": "not_found"}, status=404)

    # ---------- READ ----------
    if request.method == "GET":
        tipo = (request.GET.get("tipo") or "").strip() or None
        grupo = (request.GET.get("grupo") or "").strip() or None

        qs = Horario.objects.filter(curso=curso).order_by("dia", "hora_inicio")

        if tipo == "curso":
            qs = qs.filter(Q(tipo__isnull=True) | Q(tipo="") | Q(tipo="curso"))

        elif tipo:
            qs = qs.filter(tipo=tipo)

        if grupo:
            qs = qs.filter(grupo=grupo)


        items = []
        for h in qs:
            items.append({
                "id": h.id,
                "fecha": h.dia.isoformat(),
                "desde": h.hora_inicio.strftime("%H:%M"),
                "hasta": h.hora_fin.strftime("%H:%M"),
                "aula":  h.aula,
                "nota":  h.modulo or "",
                "tipo":  h.tipo or "",
                "grupo": h.grupo or "",
            })

        return JsonResponse({"codigo": codigo, "items": items})

    # ---------- WRITE ----------
    # простая проверка админ-токена (как у тебя в require_admin_token)
    token = request.GET.get("adm") or request.headers.get("X-MZ-Admin")
    expected = getattr(settings, "MEATZE_ADMIN_PASS", "MeatzeIT")
    if token != expected:
        return JsonResponse({"ok": False, "error": "forbidden"}, status=403)

    try:
        payload = json.loads(request.body or "{}")
    except json.JSONDecodeError:
        return JsonResponse({"ok": False, "error": "json_invalid"}, status=400)

    items = payload.get("items", [])
    if not isinstance(items, list):
        return JsonResponse({"ok": False, "error": "items_invalid"}, status=400)

    # Текущий слой (теория / практика / группа)
    current_tipo = (payload.get("tipo") or "").strip() or (request.GET.get("tipo") or "").strip()
    current_grupo = (payload.get("grupo") or "").strip() or (request.GET.get("grupo") or "").strip()

    # запасной вариант: если фронт не положил в корень, берём из первого элемента, если есть
    if not current_tipo and items:
        current_tipo = (items[0].get("tipo") or "").strip()
    if not current_grupo and items:
        current_grupo = (items[0].get("grupo") or "").strip()


    # Удаляем ТОЛЬКО текущий слой, а не всё расписание курса
    qs = Horario.objects.filter(curso=curso)

    if current_tipo == "curso" or current_tipo == "":
        qs = qs.filter(Q(tipo__isnull=True) | Q(tipo="") | Q(tipo="curso"))
    elif current_tipo:
        qs = qs.filter(tipo=current_tipo)

    if current_grupo:
        qs = qs.filter(grupo=current_grupo)

    qs.delete()


    created = 0
    for it in items:
        fecha = it.get("fecha") or it.get("dia")
        desde = it.get("desde") or it.get("inicio")
        hasta = it.get("hasta") or it.get("fin")
        if not (fecha and desde and hasta):
            continue

        try:
            d = date.fromisoformat(str(fecha))
            hi = time.fromisoformat(str(desde))
            hf = time.fromisoformat(str(hasta))
        except ValueError:
            continue

        Horario.objects.create(
            curso=curso,
            dia=d,
            hora_inicio=hi,
            hora_fin=hf,
            modulo=it.get("nota") or it.get("modulo") or "",
            aula=it.get("aula", "") or "",
            tipo=it.get("tipo") or current_tipo or "",
            grupo=it.get("grupo") or current_grupo or "",
        )
        created += 1

    return JsonResponse({"ok": True, "count": created})


@csrf_exempt
@require_admin_token
@require_http_methods(["POST"])
def admin_curso_fecha_inicio(request, codigo: str):
    """
    POST /meatze/v5/admin/curso/<codigo>/fecha_inicio?adm=...

    тело: { "fecha": "YYYY-MM-DD" }

    Обновляет поле fecha_inicio у курса с данным código.
    """
    try:
        curso = Curso.objects.get(codigo=codigo)
    except Curso.DoesNotExist:
        return JsonResponse({"ok": False, "error": "curso_not_found"}, status=404)

    try:
        payload = json.loads(request.body or "{}")
    except json.JSONDecodeError:
        return JsonResponse({"ok": False, "error": "json_invalid"}, status=400)

    fecha_str = (payload.get("fecha") or "").strip()
    if not fecha_str:
        # можно либо очищать дату, либо считать ошибкой.
        # Я сделаю "очистить", чтобы было гибко:
        curso.fecha_inicio = None
        curso.save(update_fields=["fecha_inicio"])
        return JsonResponse({"ok": True, "fecha": None})

    try:
        d = date.fromisoformat(fecha_str)
    except ValueError:
        return JsonResponse({"ok": False, "error": "fecha_invalid"}, status=400)

    curso.fecha_inicio = d
    curso.save(update_fields=["fecha_inicio"])

    return JsonResponse({"ok": True, "fecha": curso.fecha_inicio.isoformat()})
    
@api_view(["POST"])
@require_admin_token
def curso_horario_save(request, codigo):
    try:
        curso = Curso.objects.get(codigo=codigo)
    except Curso.DoesNotExist:
        return Response({"error": "not_found"}, status=404)

    items = request.data.get("items", [])
    if not isinstance(items, list):
        return Response({"error": "items_invalid"}, status=400)

    current_tipo = (request.data.get("tipo") or "").strip()
    current_grupo = (request.data.get("grupo") or "").strip()

    # запасной вариант: если фронт не положил в корень
    if not current_tipo and items:
        current_tipo = (items[0].get("tipo") or "").strip()
    if not current_grupo and items:
        current_grupo = (items[0].get("grupo") or "").strip()

    qs = Horario.objects.filter(curso=curso)

    if current_tipo == "curso" or current_tipo == "":
        qs = qs.filter(Q(tipo__isnull=True) | Q(tipo="") | Q(tipo="curso"))
    elif current_tipo:
        qs = qs.filter(tipo=current_tipo)

    if current_grupo:
        qs = qs.filter(grupo=current_grupo)

    qs.delete()


    created = 0
    for it in items:
        dia = it.get("dia") or it.get("fecha")
        inicio = it.get("inicio") or it.get("desde")
        fin = it.get("fin") or it.get("hasta")
        if not (dia and inicio and fin):
            continue

        Horario.objects.create(
            curso=curso,
            dia=dia,
            hora_inicio=inicio,
            hora_fin=fin,
            modulo=it.get("modulo") or it.get("nota") or "",
            aula=it.get("aula") or "",
            tipo=(it.get("tipo") or current_tipo or "").strip(),
            grupo=(it.get("grupo") or current_grupo or "").strip(),
        )
        created += 1

    return Response({"ok": True, "count": created})


# api/views.py
from utils.auto_horario import auto_generate_schedule
@csrf_exempt
@require_http_methods(["POST"])
@require_admin_token
def admin_auto_schedule(request, codigo):
    ...
    tipo = (payload.get("tipo") or "curso").strip()     # "curso" / "practica"
    grupo = (payload.get("grupo") or "").strip()        # для practica тут должен быть alumno_id

    # если генерируешь практику — без grupo нельзя
    if tipo == "practica" and not grupo:
        return JsonResponse({"ok": False, "error": "grupo_required_for_practica"}, status=400)

    items = auto_generate_schedule(..., grupo=(grupo or None))

    created = 0
    for it in items:
        ...
        Horario.objects.create(
            curso=curso,
            dia=d,
            hora_inicio=hi,
            hora_fin=hf,
            modulo=it.get("nota") or "",
            aula=it.get("aula") or "",
            tipo=tipo,
            grupo=grupo,
        )
        created += 1

    return JsonResponse({"ok": True, "count": created})


class AdminAuthMiddleware:
    def __init__(self, get_response):
        self.get_response = get_response

    def __call__(self, request):
        admin_token = request.headers.get("X-MZ-Admin")
        if admin_token and admin_token == settings.MZ_ADMIN_TOKEN:
            request.is_admin = True
        else:
            request.is_admin = False
        return self.get_response(request)


# api/views.py
@csrf_exempt
@require_http_methods(["POST"])
@require_admin_token
def curso_horario_bulk_delete(request, codigo):
    """
    POST /meatze/v5/curso/<codigo>/horario/bulk-delete?adm=...
    {fechas:["2025-01-10", "..."]} → удаляет эти даты у данного курса
    """
    try:
        curso = Curso.objects.get(codigo=codigo)
    except Curso.DoesNotExist:
        return JsonResponse({"error": "not_found"}, status=404)

    try:
        payload = json.loads(request.body or "{}")
    except json.JSONDecodeError:
        return JsonResponse({"ok": False, "error": "json_invalid"}, status=400)

    raw = payload.get("fechas", [])
    if not isinstance(raw, list):
        return JsonResponse({"ok": False, "error": "fechas_invalid"}, status=400)

    from datetime import date as _date
    fechas = []
    for s in raw:
        try:
            fechas.append(_date.fromisoformat(str(s)))
        except ValueError:
            continue

    qs = Horario.objects.filter(curso=curso, dia__in=fechas)

    tipo = (payload.get("tipo") or "").strip()
    grupo = (payload.get("grupo") or "").strip()

    if tipo == "curso" or tipo == "":
        qs = qs.filter(Q(tipo__isnull=True) | Q(tipo="") | Q(tipo="curso"))
    elif tipo:
        qs = qs.filter(tipo=tipo)

    if grupo:
        qs = qs.filter(grupo=grupo)

    deleted, _ = qs.delete()
    return JsonResponse({"ok": True, "deleted": deleted})




@require_admin_token
@require_http_methods(["GET"])
def curso_alumnos(request, codigo: str):
    """
    GET /meatze/v5/curso/<codigo>/alumnos?adm=...
    Возвращает список ALUMNOS (без teachers).
    """
    enrols = (
        Enrol.objects
        .filter(codigo=codigo)
        .exclude(role="teacher")          # <-- главное: учителей убираем
        .select_related("user")
        .order_by("user__first_name", "user__last_name", "user__email")
    )

    items = []
    for e in enrols:
        u = e.user
        items.append({
            "user_id": u.id,              # <-- фронт ждёт user_id
            "id": u.id,                   # <-- можно оставить для совместимости
            "email": u.email or u.username,
            "nombre": u.get_full_name() or (u.email or u.username),
            "role": e.role,
        })

    return JsonResponse({"items": items})

import zipfile
from lxml import etree

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS = {"w": W_NS}

# ===== DOCX markers (module-level) =====
COURSE_INFO_MARKERS = [
    "Nombre curso:",
    "Código curso:",
    "Año académico:",
    "Fechas impartición:",
    "Horario:",
    "Entidad:",
    "Tipo formación:",
]

import zipfile
from lxml import etree
from pathlib import Path

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS = {"w": W_NS}

COURSE_INFO_MARKERS = [
    "Nombre curso:",
    "Código curso:",
    "Año académico:",
    "Fechas impartición:",
    "Horario:",
    "Entidad:",
    "Tipo formación:",
]

def _patch_courseinfo_fonts_and_borders(
    docx_path: Path,
    markers=None,
    font_name="Trebuchet MS",
    font_size_pt=11,
    border_size=8,
    border_color="000000",
    bold_labels=True,
    row_height_cm=0.7,          # ✅ НОВОЕ
    v_align="center",           # ✅ НОВОЕ: center / top / bottom
):
    if markers is None:
        markers = COURSE_INFO_MARKERS

    half_points = str(int(font_size_pt * 2))

    def cm_to_twips(cm: float) -> str:
        # 1 inch = 2.54 cm, 1 inch = 1440 twips
        tw = int(round((cm / 2.54) * 1440))
        return str(max(tw, 1))

    row_twips = cm_to_twips(row_height_cm)

    def q(tag):
        return f"{{{W_NS}}}{tag}"

    def ensure(parent, tag, insert0=False):
        el = parent.find(f"w:{tag}", NS)
        if el is None:
            el = etree.Element(q(tag))
            if insert0 and len(parent):
                parent.insert(0, el)
            else:
                parent.append(el)
        return el

    def set_run_style(run_el, make_bold=None):
        rpr = run_el.find("w:rPr", NS)
        if rpr is None:
            rpr = etree.Element(q("rPr"))
            run_el.insert(0, rpr)

        rfonts = rpr.find("w:rFonts", NS)
        if rfonts is None:
            rfonts = etree.Element(q("rFonts"))
            rpr.insert(0, rfonts)

        rfonts.set(q("ascii"), font_name)
        rfonts.set(q("hAnsi"), font_name)
        rfonts.set(q("eastAsia"), font_name)
        rfonts.set(q("cs"), font_name)

        sz = rpr.find("w:sz", NS)
        if sz is None:
            sz = etree.Element(q("sz"))
            rpr.append(sz)
        sz.set(q("val"), half_points)

        szcs = rpr.find("w:szCs", NS)
        if szcs is None:
            szcs = etree.Element(q("szCs"))
            rpr.append(szcs)
        szcs.set(q("val"), half_points)

        if make_bold is not None:
            if make_bold:
                ensure(rpr, "b")
            else:
                b = rpr.find("w:b", NS)
                if b is not None:
                    rpr.remove(b)

    def node_text(node):
        parts = node.xpath(".//w:t/text()", namespaces=NS)
        return " ".join(" ".join(parts).split()).strip()

    def is_course_table(tbl):
        found = set()
        for tr in tbl.xpath("./w:tr", namespaces=NS):
            tcs = tr.xpath("./w:tc", namespaces=NS)
            if not tcs:
                continue
            left = node_text(tcs[0])
            for m in markers:
                if left.startswith(m):
                    found.add(m)
        return len(found) >= 3

    def set_table_borders(tbl):
        tblPr = tbl.find("w:tblPr", NS)
        if tblPr is None:
            tblPr = etree.Element(q("tblPr"))
            tbl.insert(0, tblPr)

        borders = tblPr.find("w:tblBorders", NS)
        if borders is None:
            borders = etree.Element(q("tblBorders"))
            tblPr.append(borders)

        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            el = borders.find(f"w:{edge}", NS)
            if el is None:
                el = etree.Element(q(edge))
                borders.append(el)
            el.set(q("val"), "single")
            el.set(q("sz"), str(border_size))
            el.set(q("space"), "0")
            el.set(q("color"), border_color)

    def set_row_height(tr):
        trPr = tr.find("w:trPr", NS)
        if trPr is None:
            trPr = etree.Element(q("trPr"))
            tr.insert(0, trPr)

        trH = trPr.find("w:trHeight", NS)
        if trH is None:
            trH = etree.Element(q("trHeight"))
            trPr.append(trH)

        trH.set(q("val"), row_twips)
        trH.set(q("hRule"), "exact")   # ✅ именно EXACT

    def set_cell_valign(tc):
        tcPr = tc.find("w:tcPr", NS)
        if tcPr is None:
            tcPr = etree.Element(q("tcPr"))
            tc.insert(0, tcPr)

        vA = tcPr.find("w:vAlign", NS)
        if vA is None:
            vA = etree.Element(q("vAlign"))
            tcPr.append(vA)
        vA.set(q("val"), v_align)

        # ✅ Убираем лишние отступы параграфов, чтобы “не прилипало к верху”
        for p in tc.xpath(".//w:p", namespaces=NS):
            pPr = p.find("w:pPr", NS)
            if pPr is None:
                pPr = etree.Element(q("pPr"))
                p.insert(0, pPr)

            spacing = pPr.find("w:spacing", NS)
            if spacing is None:
                spacing = etree.Element(q("spacing"))
                pPr.append(spacing)
            spacing.set(q("before"), "0")
            spacing.set(q("after"), "0")

    # ---- чтение/запись zip ----
    with zipfile.ZipFile(docx_path, "r") as zin:
        files = {name: zin.read(name) for name in zin.namelist()}

    xml = files.get("word/document.xml")
    if not xml:
        return 0, 0

    root = etree.fromstring(xml)

    hits_tables = 0
    hits_paras = 0

    # 1) TABLE курс-инфо: бордеры + шрифт + высота строк + vAlign
    for tbl in root.xpath(".//w:tbl", namespaces=NS):
        if not is_course_table(tbl):
            continue

        hits_tables += 1
        set_table_borders(tbl)

        for tr in tbl.xpath("./w:tr", namespaces=NS):
            set_row_height(tr)
            for tc in tr.xpath("./w:tc", namespaces=NS):
                set_cell_valign(tc)

        # шрифт всем runs внутри таблицы
        for r in tbl.xpath(".//w:r", namespaces=NS):
            set_run_style(r, make_bold=None)

        # жирным лейблы в первой колонке
        if bold_labels:
            for tr in tbl.xpath("./w:tr", namespaces=NS):
                tcs = tr.xpath("./w:tc", namespaces=NS)
                if not tcs:
                    continue
                left_tc = tcs[0]
                left_txt = node_text(left_tc)
                if any(left_txt.startswith(m) for m in markers):
                    for r in left_tc.xpath(".//w:r", namespaces=NS):
                        set_run_style(r, make_bold=True)

        break

    # 2) Параграфы-строки (если LO не таблицей)
    for p in root.xpath(".//w:p", namespaces=NS):
        full = node_text(p)
        if full and any(full.startswith(m) for m in markers):
            for r in p.xpath(".//w:r", namespaces=NS):
                set_run_style(r, make_bold=None)
            hits_paras += 1

    files["word/document.xml"] = etree.tostring(
        root, xml_declaration=True, encoding="UTF-8", standalone="yes"
    )

    with zipfile.ZipFile(docx_path, "w", compression=zipfile.ZIP_DEFLATED) as zout:
        for name, data in files.items():
            zout.writestr(name, data)

    return hits_tables, hits_paras


def _patch_docx_courseinfo_paragraphs(
    docx_path: Path,
    markers=None,
    font_name="Trebuchet MS",
    font_size_pt=11,
    bold_labels=True,   # ✅ ДОБАВИЛИ
):
    # ✅ если markers не передали — берём дефолт
    if markers is None:
        markers = globals().get("COURSE_INFO_MARKERS") or [
            "Nombre curso:",
            "Código curso:",
            "Año académico:",
            "Fechas impartición:",
            "Horario:",
            "Entidad:",
            "Tipo formación:",
        ]


    # ✅ на всякий случай нормализуем
    markers = [str(m).strip() for m in (markers or []) if str(m).strip()]

    half_points = str(int(font_size_pt * 2))

    def ensure_rpr(run_el):
        rpr = run_el.find("w:rPr", NS)
        if rpr is None:
            rpr = etree.Element(f"{{{W_NS}}}rPr")
            run_el.insert(0, rpr)
        return rpr

    def set_fonts(rpr):
        rfonts = rpr.find("w:rFonts", NS)
        if rfonts is None:
            rfonts = etree.Element(f"{{{W_NS}}}rFonts")
            rpr.insert(0, rfonts)
        rfonts.set(f"{{{W_NS}}}ascii", font_name)
        rfonts.set(f"{{{W_NS}}}hAnsi", font_name)
        rfonts.set(f"{{{W_NS}}}eastAsia", font_name)
        rfonts.set(f"{{{W_NS}}}cs", font_name)

        sz = rpr.find("w:sz", NS)
        if sz is None:
            sz = etree.Element(f"{{{W_NS}}}sz")
            rpr.append(sz)
        sz.set(f"{{{W_NS}}}val", half_points)

        szcs = rpr.find("w:szCs", NS)
        if szcs is None:
            szcs = etree.Element(f"{{{W_NS}}}szCs")
            rpr.append(szcs)
        szcs.set(f"{{{W_NS}}}val", half_points)

    def set_bold(rpr, on=True):
        b = rpr.find("w:b", NS)
        if b is None:
            b = etree.Element(f"{{{W_NS}}}b")
            rpr.append(b)
        b.set(f"{{{W_NS}}}val", "1" if on else "0")

    hits = 0

    with zipfile.ZipFile(docx_path, "r") as zin:
        files = {name: zin.read(name) for name in zin.namelist()}

    xml = files.get("word/document.xml")
    if not xml:
        logger.warning("DOCX XML: word/document.xml not found")
        return

    root = etree.fromstring(xml)

    # Ищем все абзацы w:p (они есть и в body и внутри w:txbxContent)
    for p in root.xpath(".//w:p", namespaces=NS):
        # склеиваем весь текст абзаца
        texts = [t.text for t in p.xpath(".//w:t", namespaces=NS) if t.text]
        full = " ".join(" ".join(texts).split()).strip()
        if not full:
            continue

        if any(full.startswith(m) for m in markers):
            # применяем стиль ко всем runs этого абзаца
            runs = p.xpath(".//w:r", namespaces=NS)
            for r in runs:
                rpr = ensure_rpr(r)
                set_fonts(rpr)

            # если хочешь: сам маркер сделать bold (аккуратно)
            # (обычно первый run и есть маркер)
            if bold_labels and runs:
                rpr0 = ensure_rpr(runs[0])
                set_bold(rpr0, True)

            hits += 1

    files["word/document.xml"] = etree.tostring(
        root, xml_declaration=True, encoding="UTF-8", standalone="yes"
    )

    with zipfile.ZipFile(docx_path, "w", compression=zipfile.ZIP_DEFLATED) as zout:
        for name, data in files.items():
            zout.writestr(name, data)

    logger.warning("DOCX XML: courseinfo paragraphs patched = %s", hits)


from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def _inject_headers_footers(docx_path: Path):
    LOGOS = getattr(settings, "MEATZE_DOCX_LOGOS", {})

    doc = Document(str(docx_path))
    section = doc.sections[0]

    # Принудительные поля секции (чтобы всё влазило и не было "узко")
    section.left_margin  = Cm(1.0)
    section.right_margin = Cm(1.0)
    section.top_margin   = Cm(1.2)
    section.bottom_margin= Cm(1.2)

    usable_width = section.page_width - section.left_margin - section.right_margin

    # ================= HEADER =================
    header = section.header
    h_table = header.add_table(rows=1, cols=2, width=usable_width)
    col_left_width  = int(usable_width * 0.5)
    col_right_width = int(usable_width * 0.5)

    h_table.columns[0].width = col_left_width
    h_table.columns[1].width = col_right_width

    cell_left, cell_right = h_table.rows[0].cells

    lanbide = LOGOS.get("lanbide")
    if lanbide and Path(lanbide).exists():
        p_l = cell_left.paragraphs[0]
        p_l.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_l.paragraph_format.space_before = Pt(0)
        p_l.paragraph_format.space_after  = Pt(0)
        p_l.add_run().add_picture(lanbide, width=Cm(2.5))

    euskadi = LOGOS.get("euskadi")
    if euskadi and Path(euskadi).exists():
        p_r = cell_right.paragraphs[0]
        p_r.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_r.paragraph_format.space_before = Pt(0)
        p_r.paragraph_format.space_after  = Pt(0)
        p_r.add_run().add_picture(euskadi, width=Cm(9.5))

    # ================= FOOTER =================
    footer = section.footer
    for p in footer.paragraphs:
        p.text = ""

    # ✅ Поднять футер (не влияет на ширину таблиц в документе)
    section.footer_distance = Cm(1.2)  # попробуй 1.0..1.5

    # ✅ ОДНА строка, без "пробела" сверху
    f_table = footer.add_table(rows=1, cols=1, width=usable_width)
    f_table.autofit = False

    eu = LOGOS.get("eu")
    if eu and Path(eu).exists():
        p = f_table.rows[0].cells[0].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        p.add_run().add_picture(eu, width=Cm(3.2))

    doc.save(str(docx_path))



@csrf_exempt
@require_http_methods(["POST"])
@require_admin_token
def export_docx_graphic(request):
    try:
        payload = json.loads(request.body or "{}")
    except json.JSONDecodeError:
        return JsonResponse({"ok": False, "error": "json_invalid"}, status=400)

    codigo = (payload.get("codigo") or "").strip().upper()
    tipo   = (payload.get("tipo") or "curso").strip().lower()
    grupo  = (payload.get("grupo") or "").strip()
    vacaciones = payload.get("vacaciones") or []

    if not codigo:
        return JsonResponse({"ok": False, "error": "codigo_required"}, status=400)

    curso = get_object_or_404(Curso, codigo=codigo)

    # ✅ ВОТ ГЛАВНОЕ: берем HTML из payload, если он есть
    html = (payload.get("html") or "").strip()

    # fallback: только если html не пришёл — тогда рендерим серверный шаблон
    if not html:
        header = build_horario_header_from_db(curso, tipo, grupo)
        mods   = build_modules_legend_from_db(curso)
        html = render_to_string("exports/calendario_docx.html", {
            "curso": curso,
            "header_main": header["main"],
            "header_exceptions": header["exceptions"],
            "modules": mods,
            "vacaciones": vacaciones,
        })

    filename = (payload.get("filename") or "calendario.docx").strip()
    if not filename.lower().endswith(".docx"):
        filename += ".docx"

    with tempfile.TemporaryDirectory() as tmpdir:
        tmpdir_path = Path(tmpdir)
        html_path = tmpdir_path / "calendario.html"
        html_path.write_text(html, encoding="utf-8")

        if not LIBREOFFICE_BIN or not Path(LIBREOFFICE_BIN).exists():
            return JsonResponse({"ok": False, "error": "soffice_not_found", "bin": LIBREOFFICE_BIN}, status=503)


        # 1) HTML -> DOCX через LibreOffice
        cmd = [
            LIBREOFFICE_BIN,
            "--headless", "--nologo",
            "--convert-to", "docx:MS Word 2007 XML",
            "--outdir", str(tmpdir_path),
            str(html_path),
        ]
        
        env = os.environ.copy()
        env["PATH"] = env.get("PATH") or "/usr/local/sbin:/usr/local/bin:/usr/sbin:/usr/bin:/sbin:/bin"

        proc = subprocess.run(
            cmd,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            cwd=tmpdir,
            timeout=120,
            env=env,
        )
        logger.info("LO stdout: %s", proc.stdout.decode(errors="ignore"))
        logger.info("LO stderr: %s", proc.stderr.decode(errors="ignore"))

        if proc.returncode != 0:
            return JsonResponse(
                {
                    "ok": False,
                    "error": "soffice_failed",
                    "stdout": proc.stdout.decode(errors="ignore"),
                    "stderr": proc.stderr.decode(errors="ignore"),
                },
                status=500,
            )

        docx_files = list(tmpdir_path.glob("*.docx"))
        if not docx_files:
            return JsonResponse({"ok": False, "error": "docx_not_generated"}, status=500)

        docx_path = docx_files[0]

        # после docx_path = docx_files[0]

        # branding как было
        try:
            _inject_headers_footers(docx_path)
        except Exception:
            logger.exception("Failed to inject headers/footers")

        # patch как было
        try:
            _patch_docx_courseinfo_paragraphs(docx_path, font_size_pt=11)
        except Exception:
            logger.exception("DOCX patch failed, continuing")

        doc = Document(str(docx_path))

        # 1) легенда модулей ИЗ БД
        try:
            mods = build_modules_legend_from_db(curso)   # <- DB
            legend = payload.get("legend") or []
            _add_modules_legend_blocks(doc, curso, legend)
        except Exception:
            logger.exception("Failed to add modules legend table")

        # 2) vacaciones / no lectivos — как раньше (payload) или тоже DB если сделаешь
        try:
            _add_legend_nonlective(doc, vacaciones)
        except Exception:
            logger.exception("Failed to add nonlective legend")

        doc.save(str(docx_path))

        # XML-бодер патч (если он “про таблицу курса”)
        try:
            _patch_courseinfo_fonts_and_borders(docx_path, font_size_pt=12, row_height_cm=0.7, v_align="center")
        except Exception:
            logger.exception("DOCX border patch failed")
        except Exception as e:
            logger.exception("Failed to add nonlective legend: %s", e)

        data = docx_path.read_bytes()

    resp = HttpResponse(
        data,
        content_type=(
            "application/vnd.openxmlformats-officedocument."
            "wordprocessingml.document"
        ),
    )
    resp["Content-Disposition"] = f"attachment; filename=\"{filename}\""
    return resp

@api_view(["POST"])
@permission_classes([AllowAny])
@authentication_classes([])
def request_pin(request):
    email = (request.data.get("email") or "").strip().lower()
    if not email:
        return Response({"message": "Email requerido"}, status=400)

    # генерируем 6 цифр
    pin = f"{random.randint(0, 999999):06d}"

    # сохраняем
    LoginPIN.objects.create(email=email, pin=pin)

    # создаём пользователя, если нет
    user, created = User.objects.get_or_create(email=email, defaults={"username": email})


    PendingRole.objects.get_or_create(
        user=user,
        email=email,
        status="pending",
        defaults={"requested_role": "unknown"},
    )


    # отправляем письмо
    send_mail(
        subject="Tu PIN de acceso a MEATZE",
        message=f"Tu código PIN es: {pin}\nVálido durante 10 minutos.",
        from_email="no-reply@meatze.eus",
        recipient_list=[email],
        fail_silently=False,
    )

    return Response({"ok": True, "message": "PIN enviado"})
