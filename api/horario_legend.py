from collections import defaultdict, Counter
from datetime import date
from django.db.models import Q
from .models import Horario, Curso  # или .models если файл рядом
from datetime import time

DEBUG_LEGEND = True

def _dbg(*a):
    if DEBUG_LEGEND:
        print("[LEGEND]", *a)
        
import re
from datetime import datetime, date

_RE_EXC_DATE = re.compile(r"\bEl\s+(\d{2}/\d{2}/\d{4})\b")

def _exc_date(line: str) -> date | None:
    if not line:
        return None
    m = _RE_EXC_DATE.search(line)
    if not m:
        return None
    try:
        return datetime.strptime(m.group(1), "%d/%m/%Y").date()
    except Exception:
        return None

def _filter_exc_by_range(lines, d_from: date | None, d_to: date | None):
    """
    Оставляет только строки-исключения, дата которых входит в [d_from, d_to].
    Строки без даты (не по формату "El dd/mm/yyyy") оставляем как есть.
    """
    if not lines:
        return []
    if not d_from or not d_to:
        return list(lines)

    out = []
    for s in lines:
        d = _exc_date(s)
        if d is None:
            out.append(s)
        else:
            if d_from <= d <= d_to:
                out.append(s)
    return out        

def _minutes(t1, t2) -> int:
    a = t1.hour * 60 + t1.minute
    b = t2.hour * 60 + t2.minute
    return b - a

def _layer_q(tipo: str, grupo: str) -> Q:
    tipo = (tipo or "curso").strip().lower()
    grupo = (grupo or "").strip()
    if tipo == "practica":
        if not grupo:
            raise ValueError("grupo es obligatorio para práctica")
        return Q(tipo="practica", grupo=grupo)
    return (Q(tipo__isnull=True) | Q(tipo="") | Q(tipo="curso"))

def _day_facts(qs):
    by_day = defaultdict(list)
    for h in qs:
        if h.dia and h.hora_inicio and h.hora_fin:
            by_day[h.dia].append(h)

    facts = {}
    for d, slots in by_day.items():
        start = min(s.hora_inicio for s in slots)
        end   = max(s.hora_fin for s in slots)
        minutes = sum(max(0, _minutes(s.hora_inicio, s.hora_fin)) for s in slots)
        facts[d] = {"start": start, "end": end, "minutes": minutes, "slots": slots}
    return facts
    
def _dominant_pattern(day_facts):
    c = Counter((v["start"], v["end"]) for v in day_facts.values())
    (base_start, base_end), _ = c.most_common(1)[0]
    return base_start, base_end
    
def _format_dmy(d: date) -> str:
    return d.strftime("%d/%m/%Y")

def build_horario_header_from_db(curso, tipo, grupo):
    layer = _layer_q(tipo, grupo)
    qs = (Horario.objects
          .filter(curso=curso).filter(layer)
          .order_by("dia", "hora_inicio", "hora_fin"))

    day_facts = _day_facts(qs)
    if not day_facts:
        return {"main": "", "exceptions": []}

    days_sorted = sorted(day_facts.keys())
    base_start, base_end = _dominant_pattern(day_facts)

    # основной период = весь диапазон (это важно!)
    main = f"Del {_format_dmy(days_sorted[0])} al {_format_dmy(days_sorted[-1])} " \
           f"{base_start.strftime('%H:%M')} – {base_end.strftime('%H:%M')}"

    exceptions = []
    for d in days_sorted:
        v = day_facts[d]
        if (v["start"], v["end"]) != (base_start, base_end):
            exceptions.append({
                "date": d,
                "line": f"El {_format_dmy(d)} el horario será de {round(v['minutes']/60,1)} horas "
                        f"(de {v['start'].strftime('%H:%M')} a {v['end'].strftime('%H:%M')})"
            })

    # ✅ если исключений много и реально есть смена режима на долго — можно дополнить режимами,
    # но твой кейс “всего один день” решается этим блоком.

    return {"main": main, "exceptions": exceptions}
   
def build_modules_legend_from_db(curso):
    qs = (Horario.objects
          .filter(curso=curso)
          .filter(Q(tipo__isnull=True) | Q(tipo="") | Q(tipo="curso"))
          .exclude(modulo__isnull=True).exclude(modulo="")
          .order_by("dia", "hora_inicio", "hora_fin"))

    by_mod = defaultdict(list)
    for h in qs:
        by_mod[h.modulo].append(h)

    out = []
    for mod_key, slots in by_mod.items():
        days = sorted({s.dia for s in slots})
        d_from, d_to = days[0], days[-1]

        by_day = defaultdict(list)
        for s in slots:
            by_day[s.dia].append(s)

        day_minutes = {}
        day_span = {}      # (start,end)
        for d, ss in by_day.items():
            mins = sum(max(0, _minutes(x.hora_inicio, x.hora_fin)) for x in ss)
            st = min(x.hora_inicio for x in ss)
            en = max(x.hora_fin for x in ss)
            day_minutes[d] = mins
            day_span[d] = (st, en)

        # ✅ типичный паттерн дня: (start,end) — самый частый
        typical_span = Counter(day_span.values()).most_common(1)[0][0]
        typical_minutes = Counter(day_minutes.values()).most_common(1)[0][0]  # оставим как доп-инфо

        exceptions = []
        for d in sorted(by_day.keys()):
            mins = day_minutes[d]
            st, en = day_span[d]

            # ✅ исключение если отличается окно (start/end) ИЛИ отличается длительность
            if (st, en) != typical_span or mins != typical_minutes:
                exceptions.append(
                    f"El {_format_dmy(d)} el horario será de {round(mins/60,1)} horas "
                    f"(de {st.strftime('%H:%M')} a {en.strftime('%H:%M')})"
                )

        out.append({
            "mod_key": mod_key,
            "from": d_from,
            "to": d_to,
            "exceptions": exceptions,
        })

    out.sort(key=lambda x: x["from"])
    return out
    
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def _set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text or "")
    run.bold = bold
    run.font.size = Pt(11)

def _set_table_borders(table, outer_sz=18, inner_sz=8, color="000000"):
    """
    outer_sz / inner_sz: Word border size units (1/8 pt). 18 ≈ 2.25pt, 8 ≈ 1pt
    """
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = tblPr.find(qn("w:tblBorders"))
    if tblBorders is None:
        tblBorders = OxmlElement("w:tblBorders")
        tblPr.append(tblBorders)

    def _border(tag, sz):
        el = tblBorders.find(qn(f"w:{tag}"))
        if el is None:
            el = OxmlElement(f"w:{tag}")
            tblBorders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)

    # внешние
    _border("top", outer_sz)
    _border("left", outer_sz)
    _border("bottom", outer_sz)
    _border("right", outer_sz)
    # внутренние
    _border("insideH", inner_sz)
    _border("insideV", inner_sz)

def _shade_row(row, fill="F3F3F3"):
    for cell in row.cells:
        tcPr = cell._tc.get_or_add_tcPr()
        shd = tcPr.find(qn("w:shd"))
        if shd is None:
            shd = OxmlElement("w:shd")
            tcPr.append(shd)
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill)

from datetime import datetime
from docx.shared import Pt

def _es_dmy(ymd: str) -> str:
    try:
        return datetime.strptime(ymd, "%Y-%m-%d").strftime("%d/%m/%Y")
    except Exception:
        return ymd or ""

def _range_es(seg: dict) -> str:
    f = _es_dmy(seg.get("from", ""))
    t = _es_dmy(seg.get("to", ""))
    if f and t and f != t:
        return f"Del {f} al {t}"
    return f or t or ""

def _add_legend_nonlective(doc, vacaciones: list):
    # Заголовок секции
    p = doc.add_paragraph("Días no lectivos / vacaciones")
    if p.runs:
        p.runs[0].bold = True
        p.runs[0].font.size = Pt(12)

    # 1) Выходные
    doc.add_paragraph("• Fines de semana (sábado y domingo) — día no lectivo")

    # 2) Сегменты vacaciones
    if not vacaciones:
        doc.add_paragraph("• (Sin días no lectivos adicionales en el rango del curso)")
        return

    for seg in vacaciones:
        rango = _range_es(seg)
        motivo = (seg.get("motivo") or "").strip()
        line = f"• {rango}" if rango else "•"
        if motivo:
            line += f" – {motivo}"
        doc.add_paragraph(line)
        
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def _shade_cell(cell, fill_hex):
    fill = (fill_hex or "").replace("#", "").strip().upper() or "FFFFFF"
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)

def _set_table_borders_nil(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = tblPr.find(qn("w:tblBorders"))
    if tblBorders is None:
        tblBorders = OxmlElement("w:tblBorders")
        tblPr.append(tblBorders)

    for tag in ("top","left","bottom","right","insideH","insideV"):
        el = tblBorders.find(qn(f"w:{tag}"))
        if el is None:
            el = OxmlElement(f"w:{tag}")
            tblBorders.append(el)
        el.set(qn("w:val"), "nil")

import re

def _hours_from_title(s: str) -> int:
    s = (s or "")
    m = re.search(r"\((\d+)\s*horas\)", s, flags=re.I)
    return int(m.group(1)) if m else 0

def _code_from_title(s: str) -> str:
    s = (s or "").strip()
    if not s:
        return ""
    # Берём первое "слово" вида MFxxxx_x / UFxxxx / MPxxxx / и т.п.
    m = re.match(r"^([A-Z]{2}\d+(?:_\d+)?)\b", s)
    return (m.group(1) if m else (s.split(" ", 1)[0].strip() if " " in s else s))

def _add_modules_legend_blocks(doc, curso, legend_items):
    """
    legend_items: список из payload.legend (у тебя legendJson):
      [{
        titulo, rango, color,
        ufs:[{titulo,rango,detalles:[]},...],
        detallesMF:[],
        horario_exceptions:[]
      }, ...]
    """

    doc.add_page_break()

    title = f"{(curso.titulo or '').strip()} — {int(getattr(curso, 'horas_total', 0) or 0)} horas"
    p = doc.add_paragraph(title)
    if p.runs:
        p.runs[0].bold = True
        p.runs[0].font.size = Pt(12)

    def _as_bullet(line: str) -> str:
        t = (line or "").strip()
        if not t:
            return ""
        return t if t.startswith("•") else f"• {t}"

    def _format_span_es(span: dict) -> str:
        if not span:
            return ""
        d1, d2 = span.get("from"), span.get("to")
        if d1 and d2:
            return f"Del {d1.strftime('%d/%m/%Y')} al {d2.strftime('%d/%m/%Y')}"
        return ""

    layer = _layer_q("curso", "")

    # ------------------------------------------------------------
    # 1) Enrich: MF/UF spans + merge exceptions, но с фильтрацией
    # ------------------------------------------------------------
    for item in (legend_items or []):
        mf_title = (item.get("titulo") or "")
        mf_code = item.get("mf_code") or _code_from_title(mf_title)

        mf_span = _span_and_exceptions_for_mfuf(curso, layer, mf_code, uf_code=None)
        if mf_span:
            item["rango"] = _format_span_es(mf_span)

            merged = (item.get("horario_exceptions") or []) + (mf_span.get("exceptions") or [])
            merged = _dedup_lines(merged)
            # ✅ ВАЖНО: выкинуть мусорные payload-строки вне MF диапазона
            merged = _filter_exc_by_range(merged, mf_span.get("from"), mf_span.get("to"))
            item["horario_exceptions"] = merged

        # UF spans
        offset_mins = 0
        for uf in (item.get("ufs") or []):
            uf_title = (uf.get("titulo") or "")
            uf_code = uf.get("uf_code") or _code_from_title(uf_title)
            uf_mins = _hours_from_title(uf_title) * 60

            uf_span = _span_and_exceptions_for_mfuf(
                curso, layer, mf_code,
                uf_code=uf_code,
                uf_target_mins=uf_mins,
                uf_offset_mins=offset_mins,
            )

            if uf_span:
                uf["rango"] = _format_span_es(uf_span)

                merged = (uf.get("detalles") or []) + (uf_span.get("exceptions") or [])
                merged = _dedup_lines(merged)
                # ✅ ВАЖНО: выкинуть мусорные payload-строки вне UF диапазона
                merged = _filter_exc_by_range(merged, uf_span.get("from"), uf_span.get("to"))
                uf["detalles"] = merged

            offset_mins += max(0, int(uf_mins or 0))

    # ------------------------------------------------------------
    # 2) Render blocks
    # ------------------------------------------------------------
    for item in (legend_items or []):
        color = (item.get("color") or "#FFFFFF")
        titulo = (item.get("titulo") or "").strip()
        rango = (item.get("rango") or "").strip()

        t = doc.add_table(rows=1, cols=1)
        _set_table_borders_nil(t)

        cell = t.rows[0].cells[0]
        _shade_cell(cell, color)
        cell.text = ""

        # MF title
        pt = cell.add_paragraph()
        run = pt.add_run(titulo)
        run.bold = True
        run.font.size = Pt(11)

        # MF rango
        if rango:
            pr = cell.add_paragraph(rango)
            if pr.runs:
                pr.runs[0].font.size = Pt(10)

        # MF details
        for line in (item.get("detallesMF") or []):
            text = _as_bullet(line)
            if text:
                pr = cell.add_paragraph(text)
                if pr.runs:
                    pr.runs[0].font.size = Pt(9.5)

        # MF exceptions
        # MF exceptions (из БД + payload)
        ufs_list = (item.get("ufs") or [])
        if not ufs_list:
            for line in (item.get("horario_exceptions") or []):
                text = _as_bullet(line)
                if text:
                    pr = cell.add_paragraph(text)
                    if pr.runs:
                        pr.runs[0].font.size = Pt(9.5)

        # UFs
        for uf in (item.get("ufs") or []):
            ut = (uf.get("titulo") or "").strip()
            ur = (uf.get("rango") or "").strip()

            if ut:
                pr = cell.add_paragraph(ut)
                if pr.runs:
                    pr.runs[0].bold = True
                    pr.runs[0].font.size = Pt(10)

            if ur:
                pr = cell.add_paragraph(ur)
                if pr.runs:
                    pr.runs[0].font.size = Pt(9.5)

            for line in (uf.get("detalles") or []):
                text = _as_bullet(line)
                if text:
                    pr = cell.add_paragraph(text)
                    if pr.runs:
                        pr.runs[0].font.size = Pt(9.5)

        doc.add_paragraph("")

def _get_layer_plan(curso: Curso, tipo: str, grupo: str):
    """
    Возвращает список (key, minutes), где key:
      - для курса: "MF0976_2|UF0349"
      - для практики: "PRACTICA|<grupo>|BLOQUE1" (или как решишь)
    """
    raw = getattr(curso, "modules", None)
    if isinstance(raw, str):
        import json
        try: raw = json.loads(raw)
        except Exception: raw = []
    if not isinstance(raw, list):
        raw = []

    plan = []
    for m in raw:
        mf = str(m.get("code") or m.get("cod") or m.get("key") or "").strip()
        # если есть ufs — дробим по UF
        ufs = m.get("ufs") or m.get("UFs") or []
        if mf and isinstance(ufs, list) and ufs:
            for uf in ufs:
                uf_code = str(uf.get("code") or uf.get("cod") or uf.get("key") or uf.get("name") or "").strip()
                horas = uf.get("hours") or uf.get("horas") or 0
                try: mins = int(round(float(horas) * 60))
                except Exception: mins = 0
                if uf_code and mins > 0:
                    plan.append((f"{mf}|{uf_code}", mins))
        else:
            # fallback на MF-уровень
            horas = m.get("hours") or m.get("horas") or 0
            try: mins = int(round(float(horas) * 60))
            except Exception: mins = 0
            if mf and mins > 0:
                plan.append((mf, mins))

    return plan
    
from django.db.models import Q
from collections import defaultdict
from datetime import date

from django.db.models import Q
import re

def _norm_exc_line(s: str) -> str:
    s = (s or "").strip()
    if not s:
        return ""

    s = s.replace("\u00A0", " ")          # NBSP -> space
    s = re.sub(r"^\s*[•\-\–]\s*", "", s)  # ✅ убрать ведущий буллет/дефис
    s = re.sub(r"\s+", " ", s)            # схлопнуть пробелы
    s = re.sub(r"(\d),(\d)", r"\1.\2", s) # 4,0 -> 4.0
    return s.strip()
def _dedup_lines(lines):
    out = []
    seen = set()
    for x in (lines or []):
        k = _norm_exc_line(x)
        if not k or k in seen:
            continue
        seen.add(k)
        out.append(k)
    return out
    
from collections import defaultdict
from django.db.models import Q
import re

def _span_and_exceptions_for_mfuf(curso, layer_q, mf_code: str, uf_code: str | None = None,
                                  uf_target_mins: int = 0, uf_offset_mins: int = 0):
    mf_code = (mf_code or "").strip()
    uf_code = (uf_code or "").strip() if uf_code else ""

    if not mf_code:
        return None

    # берём все записи курса (layer) и только текущий MF по startswith
    qs_all = (Horario.objects
              .filter(curso=curso)
              .filter(layer_q)
              .filter(Q(modulo__startswith=mf_code + " "))
              .order_by("dia", "hora_inicio", "hora_fin"))

    day_facts_all = _day_facts(qs_all)
    if not day_facts_all:
        return None

    days_all = sorted(day_facts_all.keys())

    # --- если UF не задана -> как раньше: span всего MF + исключения по доминантному окну
    if not uf_code:
        base_start, base_end = _dominant_pattern(day_facts_all)
        exceptions = []
        for d in days_all:
            v = day_facts_all[d]
            if (v["start"], v["end"]) != (base_start, base_end):
                exceptions.append(
                    f"El {_format_dmy(d)} el horario será de {round(v['minutes']/60,1)} horas "
                    f"(de {v['start'].strftime('%H:%M')} a {v['end'].strftime('%H:%M')})"
                )
        return {"from": days_all[0], "to": days_all[-1], "exceptions": exceptions}

    # --- UF задана -> режем по плану курса, а НЕ по modulo__contains
    plan = _get_layer_plan(curso, "curso", "")  # уже есть у тебя
    # вытащим только UF-части данного MF в порядке следования
    uf_plan = []
    for key, mins in plan:
        if "|" in key:
            mf, uf = key.split("|", 1)
            if mf.strip() == mf_code and uf.strip():
                uf_plan.append((uf.strip(), int(mins or 0)))
    # если плана UF нет, но нам дали target_mins (из "(44 horas)") — режем по нему
    # если плана UF нет, но нам дали target_mins + offset_mins — режем по ним
    if (not uf_plan or uf_code not in {u for u, _ in uf_plan}) and uf_target_mins > 0:
        minutes_by_day = {d: int(day_facts_all[d]["minutes"] or 0) for d in days_all}
        effective_days = [d for d in days_all if minutes_by_day.get(d, 0) > 0]
        if not effective_days:
            return None

        # 1) пропускаем uf_offset_mins
        need_skip = max(0, int(uf_offset_mins or 0))
        pos = 0
        acc = 0
        while pos < len(effective_days) and acc < need_skip:
            acc += minutes_by_day.get(effective_days[pos], 0)
            pos += 1

        if pos >= len(effective_days):
            return None

        # 2) набираем uf_target_mins
        start_day = effective_days[pos]
        need_take = max(0, int(uf_target_mins or 0))
        acc2 = 0
        end_day = start_day
        while pos < len(effective_days) and acc2 < need_take:
            d = effective_days[pos]
            acc2 += minutes_by_day.get(d, 0)
            end_day = d
            pos += 1

        # соберём факты только для UF-отрезка
        uf_days = [d for d in days_all if start_day <= d <= end_day]
        day_facts_uf = {d: day_facts_all[d] for d in uf_days if d in day_facts_all}
        if not day_facts_uf:
            return None

        base_start, base_end = _dominant_pattern(day_facts_uf)
        exceptions = []
        for d in sorted(day_facts_uf.keys()):
            v = day_facts_uf[d]
            if (v["start"], v["end"]) != (base_start, base_end):
                exceptions.append(
                    f"El {_format_dmy(d)} el horario será de {round(v['minutes']/60,1)} horas "
                    f"(de {v['start'].strftime('%H:%M')} a {v['end'].strftime('%H:%M')})"
                )

        return {"from": start_day, "to": end_day, "exceptions": exceptions}
    if not uf_plan or uf_code not in {u for u, _ in uf_plan}:
        qs_fallback = (Horario.objects
                       .filter(curso=curso)
                       .filter(layer_q)
                       .filter(Q(modulo__startswith=mf_code + " ") & Q(modulo__contains=("· " + uf_code + " ")))
                       .order_by("dia", "hora_inicio", "hora_fin"))
        day_facts = _day_facts(qs_fallback)
        if not day_facts:
            return None
        days_sorted = sorted(day_facts.keys())
        base_start, base_end = _dominant_pattern(day_facts)
        exceptions = []
        for d in days_sorted:
            v = day_facts[d]
            if (v["start"], v["end"]) != (base_start, base_end):
                exceptions.append(
                    f"El {_format_dmy(d)} el horario será de {round(v['minutes']/60,1)} horas "
                    f"(de {v['start'].strftime('%H:%M')} a {v['end'].strftime('%H:%M')})"
                )
        return {"from": days_sorted[0], "to": days_sorted[-1], "exceptions": exceptions}

    # --- основная логика нарезки UF по минутам
    minutes_by_day = {d: int(day_facts_all[d]["minutes"] or 0) for d in days_all}

    # определим границы каждой UF внутри MF (по накоплению минут)
    uf_ranges = {}
    cursor = 0
    idx = 0

    # защитимся от нулевых дней
    effective_days = [d for d in days_all if minutes_by_day.get(d, 0) > 0]
    if not effective_days:
        return None

    day_list = effective_days
    day_pos = 0

    for uf, target_mins in uf_plan:
        target_mins = max(0, int(target_mins or 0))
        if target_mins == 0:
            continue

        start_day = day_list[day_pos] if day_pos < len(day_list) else day_list[-1]
        acc = 0
        end_day = start_day

        while day_pos < len(day_list) and acc < target_mins:
            d = day_list[day_pos]
            acc += minutes_by_day.get(d, 0)
            end_day = d
            day_pos += 1

        uf_ranges[uf] = (start_day, end_day)

        if day_pos >= len(day_list):
            # дней больше нет — остальное “упрётся” в конец
            break

    if uf_code not in uf_ranges:
        return None

    d_from, d_to = uf_ranges[uf_code]

    # сформируем day_facts для выбранного UF-отрезка
    uf_days = [d for d in days_all if d_from <= d <= d_to]
    day_facts_uf = {d: day_facts_all[d] for d in uf_days if d in day_facts_all}
    if not day_facts_uf:
        return None

    base_start, base_end = _dominant_pattern(day_facts_uf)
    exceptions = []
    for d in sorted(day_facts_uf.keys()):
        v = day_facts_uf[d]
        if (v["start"], v["end"]) != (base_start, base_end):
            exceptions.append(
                f"El {_format_dmy(d)} el horario será de {round(v['minutes']/60,1)} horas "
                f"(de {v['start'].strftime('%H:%M')} a {v['end'].strftime('%H:%M')})"
            )

    return {"from": d_from, "to": d_to, "exceptions": exceptions}