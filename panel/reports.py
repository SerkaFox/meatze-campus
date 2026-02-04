# panel/reports.py
from pathlib import Path
from datetime import date as _date

from django.conf import settings
from django.contrib.auth.decorators import login_required
from django.http import FileResponse, Http404
from django.shortcuts import get_object_or_404

from api.models import Curso, Enrol
from .models import CourseTask, TaskSubmission

from .utils import (
    curso_modules_choices,
    curso_module_label,
    assign_convocatorias_auto,  # ✅ твоя функция уже в utils.py
)

from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.enum.table import WD_ALIGN_VERTICAL

# если эти helpers у тебя уже есть в views.py — вынеси в utils/reports_helpers.py,
# но для простоты можно продублировать минимум:

from docx.shared import Twips
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def _to_twips(x) -> int:
    if x is None:
        return 0
    if hasattr(x, "twips"):
        return int(x.twips)
    if hasattr(x, "emu"):
        return int(x.emu // 635)
    if isinstance(x, (int, float)):
        n = int(x)
        if n > 200000:
            return int(n // 635)
        return n
    return int(x)


def _set_table_width_twips(table, tw: int):
    tblPr = table._tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        table._tbl.insert(0, tblPr)

    for el in tblPr.findall(qn("w:tblW")):
        tblPr.remove(el)

    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:type"), "dxa")
    tblW.set(qn("w:w"), str(int(tw)))
    tblPr.append(tblW)


def _set_table_layout_fixed(table):
    tblPr = table._tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        table._tbl.insert(0, tblPr)
    for el in tblPr.findall(qn("w:tblLayout")):
        tblPr.remove(el)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)


def _set_table_indent_zero(table):
    tblPr = table._tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        table._tbl.insert(0, tblPr)
    for el in tblPr.findall(qn("w:tblInd")):
        tblPr.remove(el)
    tblInd = OxmlElement("w:tblInd")
    tblInd.set(qn("w:w"), "0")
    tblInd.set(qn("w:type"), "dxa")
    tblPr.append(tblInd)


def _apply_col_widths_to_all_rows(table, widths_twips: list[int]):
    for i, w in enumerate(widths_twips):
        try:
            table.columns[i].width = Twips(int(w))
        except Exception:
            pass
    for row in table.rows:
        for i, w in enumerate(widths_twips):
            row.cells[i].width = Twips(int(w))


def _apply_cell_style(cell, *, font_pt=10, bold=False, align=None):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    # гарантируем run
    if not p.runs:
        p.add_run("")
    for r in p.runs:
        r.font.size = Pt(font_pt)
        r.bold = bool(bold)


@login_required
def teacher_anexo_vi_report_doc(request, codigo: str):
    curso = get_object_or_404(Curso, codigo=codigo)

    # доступ только teacher этого курса
    is_teacher = Enrol.objects.filter(user=request.user, codigo=curso.codigo, role="teacher").exists()
    if not is_teacher:
        raise Http404

    fecha = _date.today()

    # ✅ фильтр по модулю
    selected_mod = (request.GET.get("mod") or "").strip()

    # валидные модули
    valid_mods = {k for k, _ in curso_modules_choices(curso)}
    if selected_mod and selected_mod not in valid_mods:
        # если мусор — лучше 404, чтобы не скачивали "пустоту"
        raise Http404

    # --- студенты курса ---
    enrols = (
        Enrol.objects
        .filter(codigo=curso.codigo, role="student")
        .select_related("user")
        .order_by("created_at", "id")
    )
    student_ids = [e.user_id for e in enrols]

    # --- финальные экзамены (по модулю) ---
    exams_qs = CourseTask.objects.filter(curso=curso, is_final_exam=True)

    if selected_mod:
        exams_qs = exams_qs.filter(module_key=selected_mod)

    exams = list(exams_qs.order_by("module_key", "created_at", "id"))

    # группировка: module_key -> [tasks...]
    by_mod = {}
    for t in exams:
        mk = (t.module_key or "").strip()
        if not mk:
            mk = "(sin módulo)"
        by_mod.setdefault(mk, []).append(t)

    # slots 1ª/2ª
    mod_slots = {}
    for mk, tasks_sorted in by_mod.items():
        mod_slots[mk] = assign_convocatorias_auto(tasks_sorted)

    # список модулей в отчёте
    mod_list = list(mod_slots.keys())

    # --- оценки ---
    task_ids = [t.id for t in exams]
    subs = list(
        TaskSubmission.objects
        .filter(task_id__in=task_ids,
                alumno_id__in=student_ids,
                status=TaskSubmission.STATUS_GRADED)
        .order_by("task_id", "alumno_id", "-graded_at", "-id")
        .values("task_id", "alumno_id", "grade")
    )

    grade_map = {}
    for s in subs:
        key = (s["task_id"], s["alumno_id"])
        if key not in grade_map:
            grade_map[key] = s["grade"]

    # alumnos_rows
    alumnos_rows = []
    for e in enrols:
        u = e.user
        profile = getattr(u, "profile", None)
        last1 = (getattr(profile, "last_name1", "") or "").strip()
        last2 = (getattr(profile, "last_name2", "") or "").strip()
        first = (getattr(profile, "first_name", "") or u.first_name or "").strip()
        apellidos = " ".join([x for x in [last1, last2] if x]).strip()
        nombre = f"{apellidos}, {first}".strip().strip(",") or (u.get_full_name() or u.email or u.username or f"ID {u.pk}")

        alumnos_rows.append({"id": u.pk, "nombre": nombre, "email": (u.email or "").strip()})

    out_dir = Path(getattr(settings, "MEDIA_ROOT", "/tmp")) / "tmp_docs"
    out_dir.mkdir(parents=True, exist_ok=True)

    mod_suffix = "ALL" if not selected_mod else "".join(ch for ch in selected_mod if ch.isalnum() or ch in ("-", "_"))[:40]
    out_path = out_dir / f"ANEXO_VI_{curso.codigo}_{mod_suffix}_{fecha.strftime('%Y%m%d')}.docx"

    build_anexo_vi_docx(
        curso=curso,
        fecha=fecha,
        alumnos_rows=alumnos_rows,
        mod_list=mod_list,
        mod_slots=mod_slots,
        grade_map=grade_map,
        out_path=out_path,
    )

    # если ты хочешь твою шапку/футер с логотипами:
    from .views import inject_headers_footers_original  # или откуда у тебя
    inject_headers_footers_original(out_path)

    return FileResponse(open(out_path, "rb"), as_attachment=True, filename=out_path.name)


def build_anexo_vi_docx(*, curso, fecha, alumnos_rows, mod_list, mod_slots, grade_map, out_path: Path):
    doc = Document()
    section = doc.sections[0]
    section.left_margin = Cm(1.0)
    section.right_margin = Cm(1.0)
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)

    usable_width = section.page_width - section.left_margin - section.right_margin
    tw = _to_twips(usable_width)

    # заголовок
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("ANEXO VI – Evaluación (Exámenes finales)")
    r.bold = True
    r.font.size = Pt(13)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.add_run(
        f"Curso: {(getattr(curso, 'titulo', '') or '').strip()}  ·  Código: {curso.codigo}  ·  Fecha: {fecha.strftime('%d/%m/%Y')}"
    ).font.size = Pt(10)

    doc.add_paragraph("")

    cols = 2 + (len(mod_list) * 2)
    tbl = doc.add_table(rows=2, cols=cols)
    tbl.style = "Table Grid"
    tbl.autofit = False
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    _set_table_indent_zero(tbl)
    _set_table_layout_fixed(tbl)
    _set_table_width_twips(tbl, tw)

    w_num = int(tw * 0.06)
    w_name = int(tw * 0.24)
    rest = tw - w_num - w_name
    per = int(rest / max(1, (len(mod_list) * 2)))

    widths = [w_num, w_name] + [per] * (len(mod_list) * 2)
    _apply_col_widths_to_all_rows(tbl, widths)

    row0 = tbl.rows[0].cells
    row1 = tbl.rows[1].cells

    row0[0].text = "Nº"
    row0[1].text = "ALUMN@"
    _apply_cell_style(row0[0], font_pt=9, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _apply_cell_style(row0[1], font_pt=9, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    row1[0].text = ""
    row1[1].text = ""
    _apply_cell_style(row1[0], font_pt=9, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _apply_cell_style(row1[1], font_pt=9, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    c = 2
    for mk in mod_list:
        label = curso_module_label(curso, mk)  # ✅ красиво
        a = tbl.cell(0, c)
        b = tbl.cell(0, c + 1)
        a.merge(b)
        a.text = (label or mk or "").strip()
        _apply_cell_style(a, font_pt=8, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

        tbl.cell(1, c).text = "1ª"
        tbl.cell(1, c + 1).text = "2ª"
        _apply_cell_style(tbl.cell(1, c), font_pt=9, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        _apply_cell_style(tbl.cell(1, c + 1), font_pt=9, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        c += 2

    for rr in (tbl.rows[0], tbl.rows[1]):
        rr.height = Cm(0.7)
        rr.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

    def _fmt(g):
        if g is None:
            return ""
        try:
            x = float(g)
            return str(int(x)) if abs(x - int(x)) < 1e-9 else str(x).replace(".", ",")
        except Exception:
            return ""

    for idx, a in enumerate(alumnos_rows, start=1):
        r = tbl.add_row()
        _apply_col_widths_to_all_rows(tbl, widths)
        r.height = Cm(0.65)
        r.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

        r.cells[0].text = str(idx)
        r.cells[1].text = (a.get("nombre") or "").strip().upper()
        _apply_cell_style(r.cells[0], font_pt=9, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER)
        _apply_cell_style(r.cells[1], font_pt=9, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT)

        sid = a["id"]
        col = 2
        for mk in mod_list:
            slots = mod_slots.get(mk) or {1: None, 2: None}
            t1 = slots.get(1)
            t2 = slots.get(2)
            g1 = grade_map.get((t1.id, sid)) if t1 else None
            g2 = grade_map.get((t2.id, sid)) if t2 else None

            r.cells[col].text = _fmt(g1)
            r.cells[col + 1].text = _fmt(g2)
            _apply_cell_style(r.cells[col], font_pt=9, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER)
            _apply_cell_style(r.cells[col + 1], font_pt=9, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER)
            col += 2

    doc.save(str(out_path))
