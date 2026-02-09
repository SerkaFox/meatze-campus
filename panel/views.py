from .decorators import require_profile_complete
from .models import CursoFile, CursoFolder, MaterialDownload, MaterialReceipt, CursoPhysicalConfig, AttendanceSession, CourseTask, TaskSubmission, PublicShareLink, CursoPhysicalItem
from .utils import curso_modules_choices, curso_module_label, extract_convocatoria, resolve_convocatoria 
from .views_attendance import _current_slot
from api.models import Curso, Enrol, Horario
from datetime import date as _date
from datetime import timedelta
from django.conf import settings
from django.contrib.auth import get_user_model
from django.contrib.auth.decorators import login_required
from django.core.files.base import ContentFile
from django.db import IntegrityError, transaction, models
from django.db.models import Q, Count, Prefetch, F, Avg
from django.http import HttpResponse, JsonResponse, HttpResponseBadRequest, Http404
from django.shortcuts import render, get_object_or_404, redirect
from django.utils import timezone
from django.utils.crypto import get_random_string
from django.utils.text import slugify
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.http import require_http_methods
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, Emu, Twips
from pathlib import Path
from tempfile import TemporaryDirectory
from urllib.parse import urlparse
import random, re, requests
from collections import defaultdict
from django.template.loader import render_to_string
from .materiales_fs import normalize_path, create_folder

try:
    from api.models import UserProfile
except ImportError:
    UserProfile = None
import logging
logger = logging.getLogger(__name__)
User = get_user_model()


TEACHER_MODULES = [
    {"slug": "info",       "label": "Información"},
    {"slug": "materiales", "label": "Materiales"},
    {"slug": "tareas", "label": "Tareas"}, 
    {"slug": "calendario", "label": "Calendario"},
    {"slug": "alumnos",    "label": "Alumnos"},
    {"slug": "ia",         "label": "IA"},
    {"slug": "chat",       "label": "Chat"},
]

STUDENT_MODULES = [
    {"slug": "info",       "label": "Información"},
    {"slug": "materiales", "label": "Materiales"},
    {"slug": "tareas", "label": "Tareas"}, 
    {"slug": "calendario", "label": "Calendario"},
    {"slug": "chat",       "label": "Chat"},
]

PHYSICAL_ITEMS = [
    {"key": "notebook", "label": "Cuaderno"},
    {"key": "pen", "label": "Bolígrafo"},
    {"key": "printed_docs", "label": "Documentación impresa"},
    {"key": "usb", "label": "Memoria USB"},
]

MAX_URL_SIZE = 25 * 1024 * 1024  # 25MB
URL_TIMEOUT = (5, 25)            # connect/read

def _safe_filename(name: str) -> str:
    name = (name or "file").strip()
    name = re.sub(r'[^A-Za-z0-9._-]+', '_', name).strip("._")
    return (name[:140] or "file")

def _filename_from_url(url: str) -> str:
    p = urlparse(url)
    base = (p.path or "").split("/")[-1]
    return _safe_filename(base or "file")


def _normalize_path(p: str) -> str:
    p = (p or "").strip().replace("\\", "/").strip("/")
    p = re.sub(r"/{2,}", "/", p)
    return p

def _ensure_module_folders(curso, tmp_mods, user):
    """
    tmp_mods — это твой tmp (MF/UF структура)
    Создаём папки только для MF/UF имён (как у тебя в фильтрах mod_groups).
    """
    wanted = []
    for m in tmp_mods:
        name = (m.get("name") or "").strip()
        if not name:
            continue
        # MF и UF делаем папками (можно ограничить MF-only — но ты просил по названию модулей)
        wanted.append(_normalize_path(name))

    for path in wanted:
        if not path:
            continue
        CursoFolder.objects.get_or_create(
            curso=curso,
            path=path,
            defaults={"title": path.split("/")[-1], "created_by": user, "is_locked": True}
        )


def _download_url(url: str):
    p = urlparse(url)
    if p.scheme not in ("http", "https"):
        raise ValueError("URL debe empezar por http(s)://")

    r = requests.get(url, stream=True, timeout=URL_TIMEOUT, allow_redirects=True)
    r.raise_for_status()

    cd = r.headers.get("Content-Disposition", "") or ""
    m = re.search(r'filename\*?=(?:UTF-8\'\')?"?([^";]+)"?', cd, re.IGNORECASE)
    filename = _safe_filename(m.group(1)) if m else _filename_from_url(r.url)

    total = 0
    chunks = []
    for chunk in r.iter_content(chunk_size=1024 * 128):
        if not chunk:
            continue
        total += len(chunk)
        if total > MAX_URL_SIZE:
            raise ValueError("Archivo demasiado grande")
        chunks.append(chunk)

    return filename, b"".join(chunks)


@login_required
def alumnos_status(request, codigo: str):
    curso = get_object_or_404(Curso, codigo=codigo)

    is_teacher = Enrol.objects.filter(user=request.user, codigo=curso.codigo, role="teacher").exists()

    # ✅ реальный список физ. предметов из конфигурации курса
    physical_keys = list(
        CursoPhysicalItem.objects.filter(curso=curso, is_enabled=True)
        .values_list("key", flat=True)
    )

    # какие файлы видны ученикам
    visible_files_qs = CursoFile.objects.filter(curso=curso).filter(
        Q(tipo=CursoFile.TIPO_ALUMNOS) |
        (Q(share_alumnos=True) & Q(tipo__in=[CursoFile.TIPO_DOCENTES, CursoFile.TIPO_PRIVADO]))
    ).only("id", "title", "file", "module_key", "ext", "size")

    visible_ids = list(visible_files_qs.values_list("id", flat=True))
    total_visible = len(visible_ids)

    # ✅ STUDENT MODE (без нового эндпоинта)
    if not is_teacher:
        # проверим доступ к курсу (ученик должен быть зачислен)
        is_student = Enrol.objects.filter(user=request.user, codigo=curso.codigo).exclude(role="teacher").exists()
        if not is_student:
            return JsonResponse({"message": "Acceso denegado"}, status=403)

        my_d = MaterialDownload.objects.filter(alumno=request.user, file_id__in=visible_ids).count()
        my_r = MaterialReceipt.objects.filter(curso=curso, alumno=request.user, item_key__in=physical_keys).count()

        return JsonResponse({
            "role": "student",
            "curso_codigo": curso.codigo,
            "total_files": total_visible,
            "total_phys": len(physical_keys),
            "my_d": int(my_d),
            "my_r": int(my_r),
        })


    # ----- агрегаты (счетчики) -----
    dl_counts = dict(
        MaterialDownload.objects.filter(file_id__in=visible_ids)
        .values("alumno_id").annotate(c=Count("id"))
        .values_list("alumno_id", "c")
    )
    rc_counts = dict(
        MaterialReceipt.objects.filter(curso=curso, item_key__in=physical_keys)
        .values("alumno_id").annotate(c=Count("id"))
        .values_list("alumno_id", "c")
    )

    # ----- детали для tooltip -----
    # downloads: alumno -> [file_id,...]
    dl_rows = (
        MaterialDownload.objects
        .filter(file_id__in=visible_ids)
        .values("alumno_id", "file_id")
    )

    file_map = {f.id: (f.title or f.filename) for f in visible_files_qs}

    dl_detail = {}
    for row in dl_rows:
        aid = row["alumno_id"]
        fid = row["file_id"]
        dl_detail.setdefault(aid, []).append(file_map.get(fid, f"#{fid}"))

    # receipts: alumno -> [item_label,...]
    rc_rows = (
        MaterialReceipt.objects
        .filter(curso=curso, item_key__in=physical_keys)
        .values("alumno_id", "item_label", "item_key")
    )
    # fallback label по ключу
    label_by_key = dict(
    CursoPhysicalItem.objects.filter(curso=curso)
    .values_list("key", "label")
    )

    rc_detail = {}
    for row in rc_rows:
        aid = row["alumno_id"]
        label = (row.get("item_label") or "").strip() or label_by_key.get(row["item_key"], row["item_key"])
        rc_detail.setdefault(aid, []).append(label)

    # все alumno_id, у кого есть хоть что-то
    all_ids = set(dl_counts.keys()) | set(rc_counts.keys()) | set(dl_detail.keys()) | set(rc_detail.keys())
    slot = _current_slot(curso)
    presence = {}

    if slot:
        day_start = timezone.now().replace(hour=0, minute=0, second=0, microsecond=0)

        # берём ВСЕ сессии этого слота за сегодня
        qs = (AttendanceSession.objects
              .filter(curso_codigo=curso.codigo, horario_id=slot.id, started_at__gte=day_start)
              .exclude(status=AttendanceSession.STATUS_ENDED)
              # ключевое: сначала самые “живые”
              .order_by("-last_heartbeat_at", "-last_seen_at", "-started_at", "-id")
              .values("id","user_id","status","active_sec","active_confirmed_sec",
                      "started_at","confirmed_at","last_heartbeat_at"))

        for r in qs:
            uid = str(r["user_id"])
            if uid in presence:
                continue  # уже взяли самую свежую для этого user

            presence[uid] = {
                "id": r["id"],
                "status": r["status"],
                "active_sec": int(r["active_sec"] or 0),
                "confirmed": bool(r["confirmed_at"]),
                "last": r["last_heartbeat_at"].isoformat() if r["last_heartbeat_at"] else None,
                "started": r["started_at"].isoformat() if r["started_at"] else None,
            }

    students_total = Enrol.objects.filter(codigo=curso.codigo, role="student").count()


    return JsonResponse({
        "total_files": total_visible,
        "total_phys": len(physical_keys),
        "students_total": students_total,

        "items": {
            str(aid): {
                "d": int(dl_counts.get(aid, 0)),
                "r": int(rc_counts.get(aid, 0)),
                "dl": dl_detail.get(aid, []),   # ✅ названия скачанных файлов
                "rc": rc_detail.get(aid, []),   # ✅ названия подтверждённых предметов
            }
            for aid in all_ids
        },
        "slot": {"id": slot.id, "dia": slot.dia.isoformat(), "desde": slot.hora_inicio.strftime("%H:%M"), "hasta": slot.hora_fin.strftime("%H:%M")} if slot else None,
        "presence": presence
       
    })


@login_required
@require_profile_complete
def course_panel(request, codigo):
    curso = get_object_or_404(Curso, codigo=codigo)
    user = request.user
    students = []
    teachers = []

    is_teacher = Enrol.objects.filter(
        user=user,
        codigo=curso.codigo,
        role="teacher",
    ).exists()

    modules_cfg = TEACHER_MODULES if is_teacher else STUDENT_MODULES
    active_tab = request.GET.get("tab", "info")

    # --- варианты аудитории для вкладки "Materiales" ---
    if is_teacher:
        audience_options = [
            ("alumnos", "Para alumnos"),
            ("docentes", "Docentes"),
            ("mis", "Privados (míos)"),
        ]
    else:
        # у студента одна кнопка, просто текст красивый
        audience_options = [
            ("alumnos", "Materiales del curso"),
        ]

    # --- модули курса для вкладки INFO ---
    raw_mods = curso.modules or []

    def _get_name(m):
        return (m.get("name") or m.get("nombre") or m.get("titulo") or "").strip()

    def _get_code(m):
        return (m.get("code") or m.get("codigo") or m.get("clave") or "").strip()

    def _get_direct_hours(m):
        h = m.get("hours") or m.get("horas") or 0
        try:
            return float(h)
        except (TypeError, ValueError):
            return 0.0

    def _guess_kind(m):
        """
        Пытаемся понять MF/UF по коду или первому слову.
        """
        code = _get_code(m)
        name = _get_name(m)
        token = (code or name).split()[0].upper()
        if token.startswith("MF"):
            return "MF"
        if token.startswith("UF"):
            return "UF"
        return ""

    # временная структура с привязкой UF -> MF
    tmp = []
    current_mf_idx = None
    parent_for_child = {}

    for m in raw_mods:
        if not isinstance(m, dict):
            continue

        name = _get_name(m)
        code = _get_code(m)
        kind = _guess_kind(m)
        hours_direct = _get_direct_hours(m)

        idx = len(tmp)
        tmp.append({
            "name": name,
            "code": code,
            "kind": kind,          # "MF" / "UF" / ""
            "hours_direct": hours_direct,
            "hours": hours_direct,  # финальное значение заполним ниже
            "children": [],
        })

        if kind == "MF":
            current_mf_idx = idx
        elif kind == "UF" and current_mf_idx is not None:
            tmp[current_mf_idx]["children"].append(idx)
            parent_for_child[idx] = current_mf_idx
        else:
            current_mf_idx = None

    # 1) для MF с 0 часами — суммируем их UF
    for i, mod in enumerate(tmp):
        if mod["kind"] == "MF":
            if mod["hours_direct"] > 0:
                continue
            children = mod.get("children") or []
            agg = 0.0
            for c_idx in children:
                agg += tmp[c_idx]["hours_direct"]
            mod["hours"] = agg
        else:
            mod["hours"] = mod["hours_direct"]

    # 2) считаем общую длительность, не удваивая UF
    total_hours = 0.0
    for i, mod in enumerate(tmp):
        if mod["kind"] == "MF":
            total_hours += mod["hours"]
        elif mod["kind"] == "UF":
            # добавляем только UF, которые ни к одному MF не привязаны
            if i not in parent_for_child:
                total_hours += mod["hours"]
        else:
            total_hours += mod["hours"]

    # 3) простой список для шаблона
    mods = [{"name": m["name"], "hours": m["hours"]} for m in tmp]

    # базовый контекст — один раз
    context = {
        "curso": curso,
        "mods": mods,
        "total_hours": total_hours,   # <<< добавили
        "is_teacher": is_teacher,
        "modules_cfg": modules_cfg,
        "active_tab": active_tab,
        "audience_options": audience_options,
    }
    context["IS_TEACHER"] = request.user.groups.filter(name="Docente").exists()
    context["USER_ID"] = request.user.id

    # ───────────── ВКЛАДКА "MATERIALES" ─────────────
    if active_tab == "materiales":
        audience = request.GET.get("aud", "alumnos")  # alumnos/docentes/mis
        selected_mod = ""
        sort = (request.GET.get("sort") or "new").strip()  # new/old/az/za
        if sort not in {"new", "old", "az", "za"}:
            sort = "new"
        context["sort"] = sort
        if is_teacher:
            _ensure_module_folders(curso, tmp, request.user)

        current_path = _normalize_path(request.GET.get("p") or "")
        context["current_path"] = current_path
        folders_qs = CursoFolder.objects.filter(curso=curso, is_deleted=False)

        # показываем только “детей” текущей папки:
        prefix = current_path + "/" if current_path else ""
        level = prefix.count("/") + (1 if prefix else 0)

        # простой способ: фильтруем по startswith и берём только следующий уровень
        cand = folders_qs.filter(path__startswith=prefix).values_list("path", "title", "is_locked")
        children = []
        seen = set()
        for pth, title, locked in cand:
            rest = pth[len(prefix):]
            if not rest:
                continue
            first = rest.split("/", 1)[0]
            child_path = (prefix + first) if prefix else first
            if child_path in seen:
                continue
            seen.add(child_path)
            children.append({"path": child_path, "name": first, "is_locked": locked})

        context["folders"] = sorted(children, key=lambda x: (0 if x["is_locked"] else 1, x["name"].lower()))


        if not is_teacher:
            audience = "alumnos"

        # --- форматированный размер файла ---
        def _fmt_size(num):
            try:
                n = float(num or 0)
            except (TypeError, ValueError):
                n = 0.0
            for unit in ("B", "KB", "MB", "GB", "TB"):
                if n < 1024.0 or unit == "TB":
                    if unit == "B":
                        return f"{int(n)} B"
                    return f"{n:.1f} {unit}"
                n /= 1024.0
            return f"{n} B"

        # ===== Physical config (teacher chooses what exists) =====
        items_qs = CursoPhysicalItem.objects.filter(curso=curso).order_by("label", "id")
        physical_items_all = list(items_qs)
        physical_items_enabled = [x for x in physical_items_all if x.is_enabled]

        context.update({
            "physical_items_all": physical_items_all,         # teacher list
            "physical_items_enabled": physical_items_enabled, # student list
        })

        # ===== Teacher: save physical config =====
        if request.method == "POST" and is_teacher and request.POST.get("action") == "physical_item_add":
            label = (request.POST.get("label") or "").strip()
            if label:
                base = slugify(label)[:40] or "item"
                key = base
                n = 2
                while CursoPhysicalItem.objects.filter(curso=curso, key=key).exists():
                    key = f"{base}-{n}"
                    n += 1

                CursoPhysicalItem.objects.create(
                    curso=curso,
                    key=key,
                    label=label,
                    is_enabled=True,
                    order=0
                )
            return redirect(f"{request.path}?tab=materiales&aud={audience}&sort={sort}")


        if request.method == "POST" and is_teacher and request.POST.get("action") == "physical_item_update":
            item_id = request.POST.get("item_id")
            it = CursoPhysicalItem.objects.filter(pk=item_id, curso=curso).first()
            if it:
                it.label = (request.POST.get("label") or it.label).strip()
                it.is_enabled = (request.POST.get("is_enabled") == "on")
                it.save()
            return redirect(f"{request.path}?tab=materiales&aud={audience}&sort={sort}")

        if request.method == "POST" and is_teacher and request.POST.get("action") == "physical_item_delete":
            item_id = request.POST.get("item_id")
            CursoPhysicalItem.objects.filter(pk=item_id, curso=curso).delete()
            return redirect(f"{request.path}?tab=materiales&aud={audience}&sort={sort}")

        if request.method == "POST" and (not is_teacher) and request.POST.get("action") == "receipt_save":
            enabled_items = list(
                CursoPhysicalItem.objects.filter(curso=curso, is_enabled=True)
                .values("key", "label")
            )

            existing = set(
                MaterialReceipt.objects.filter(curso=curso, alumno=user)
                .values_list("item_key", flat=True)
            )

            for it in enabled_items:
                k = it["key"]
                if k in existing:
                    continue
                if request.POST.get(f"receipt_{k}") == "on":
                    MaterialReceipt.objects.create(
                        curso=curso,
                        alumno=user,
                        item_key=k,
                        item_label=it["label"],
                    )

            return redirect(f"{request.path}?tab=materiales&aud={audience}&sort={sort}")
        if not is_teacher:
            receipt_keys = set(MaterialReceipt.objects.filter(curso=curso, alumno=user).values_list("item_key", flat=True))
            enabled_keys = set(CursoPhysicalItem.objects.filter(curso=curso, is_enabled=True).values_list("key", flat=True))
            context.update({
                "receipt_keys": receipt_keys,
                "receipt_locked_keys": set(receipt_keys),
                "receipt_all_done": (len(enabled_keys) > 0 and receipt_keys.issuperset(enabled_keys)),
            })


        # ===== “Así lo ven los alumnos” (teacher preview) =====
        if is_teacher:
            # ✅ ВАЖНО: preview всегда без фильтра по модулю
            student_visible_qs = CursoFile.objects.filter(curso=curso).filter(
                Q(tipo=CursoFile.TIPO_ALUMNOS) |
                (Q(share_alumnos=True) & Q(tipo__in=[CursoFile.TIPO_DOCENTES, CursoFile.TIPO_PRIVADO]))
            )

            student_visible_files = list(student_visible_qs.order_by("-created_at")[:60])
            for f in student_visible_files:
                f.fmt_size = _fmt_size(getattr(f, "size", 0))

            context["student_visible_files"] = student_visible_files

        # ===== base queryset for this audience =====
        files_qs = CursoFile.objects.filter(curso=curso)

        if is_teacher:
            if audience == "alumnos":
                files_qs = files_qs.filter(tipo=CursoFile.TIPO_ALUMNOS)
            elif audience == "docentes":
                files_qs = files_qs.filter(tipo=CursoFile.TIPO_DOCENTES)

            elif audience == "mis":
                files_qs = files_qs.filter(tipo=CursoFile.TIPO_PRIVADO, uploaded_by=user)
        else:
            files_qs = files_qs.filter(
                Q(tipo=CursoFile.TIPO_ALUMNOS) |
                Q(tipo__in=[CursoFile.TIPO_DOCENTES, CursoFile.TIPO_PRIVADO], share_alumnos=True)
            )
            
        # ===== TREE VIEW (folders + files) =====
        view = (request.GET.get("view") or "tree").strip()
        if view not in {"tree", "cards"}:
            view = "tree"
        context["view"] = view

        if view == "tree":
            # 1) все файлы по аудитории (без folder_path фильтра)
            tree_qs = CursoFile.objects.filter(curso=curso)

            if is_teacher:
                if audience == "alumnos":
                    tree_qs = tree_qs.filter(tipo=CursoFile.TIPO_ALUMNOS)
                elif audience == "docentes":
                    tree_qs = tree_qs.filter(tipo=CursoFile.TIPO_DOCENTES)
                elif audience == "mis":
                    tree_qs = tree_qs.filter(tipo=CursoFile.TIPO_PRIVADO, uploaded_by=user)
            else:
                tree_qs = tree_qs.filter(
                    Q(tipo=CursoFile.TIPO_ALUMNOS) |
                    Q(tipo__in=[CursoFile.TIPO_DOCENTES, CursoFile.TIPO_PRIVADO], share_alumnos=True)
                )

            # сорт внутри папки
            tree_qs = tree_qs.order_by("-created_at", "-id")
            tree_files = list(tree_qs.order_by("-created_at", "-id"))
            # downloaded marker for students (tree)
            if not is_teacher:
                dl_ids = set(
                    MaterialDownload.objects.filter(alumno=user, file_id__in=[f.id for f in tree_files])
                    .values_list("file_id", flat=True)
                )
                for f in tree_files:
                    f.is_downloaded = (f.id in dl_ids)

            # 2) папки
            folders_all = list(
                CursoFolder.objects.filter(curso=curso, is_deleted=False)
                .values("path", "title", "is_locked")
            )

            folder_by_path = {x["path"]: x for x in folders_all}

            # 3) group files by folder_path

            files_by_folder = defaultdict(list)
            root_files = []
            for f in tree_files:
                fp = (f.folder_path or "").strip()
                if fp:
                    files_by_folder[fp].append(f)
                else:
                    root_files.append(f)

            # --- helpers ---
            def parent_path(p: str) -> str:
                p = (p or "").strip().strip("/")
                if not p: return ""
                if "/" not in p: return ""
                return p.rsplit("/", 1)[0]

            def leaf_name(p: str) -> str:
                p = (p or "").strip().strip("/")
                return p.split("/")[-1] if p else ""

            # 1) Убедимся что в дереве есть все пути папок, которые упоминаются файлами
            all_paths = set(folder_by_path.keys()) | set(files_by_folder.keys())

            # 2) Для каждого пути добавим всех родителей, чтобы дерево не рвалось
            expanded = set()
            for p in list(all_paths):
                cur = p
                while cur:
                    expanded.add(cur)
                    cur = parent_path(cur)
            all_paths = expanded

            # 3) children map: parent -> [child1, child2]
            children = defaultdict(list)
            for p in all_paths:
                children[parent_path(p)].append(p)

            # 4) сортировка детей: locked вверх, затем A-Z
            def node_sort_key(path):
                meta = folder_by_path.get(path) or {}
                locked = bool(meta.get("is_locked"))
                nm = (meta.get("title") or leaf_name(path) or path).lower()
                return (0 if locked else 1, nm)

            for par in list(children.keys()):
                children[par] = sorted(children[par], key=node_sort_key)

            # 5) Рекурсивно строим nodes
            def build_node(path):
                meta = folder_by_path.get(path) or {}
                name = (meta.get("title") or leaf_name(path) or path)
                locked = bool(meta.get("is_locked"))
                node_files = files_by_folder.get(path, [])

                kids = [build_node(ch) for ch in children.get(path, [])]

                count_direct = len(node_files)
                count_total = count_direct + sum(int(k.get("count_total", 0)) for k in kids)

                return {
                    "path": path,
                    "name": name,
                    "is_locked": locked,
                    "count_direct": count_direct,
                    "count_total": count_total,
                    "files": node_files,
                    "children": kids,
                }

            tree_nodes = [build_node(ch) for ch in children.get("", [])]

            context["tree_root_files"] = root_files
            context["tree_nodes"] = tree_nodes
            
        files_qs = files_qs.filter(folder_path=current_path)

        # --- teacher POST ops (upload/delete/share/copy) ---
        if request.method == "POST" and is_teacher:
            action = (request.POST.get("action") or "").strip()
            
            
            if action == "folder_delete":
                path = _normalize_path(request.POST.get("path") or "")
                mode = (request.POST.get("mode") or "check").strip()

                f = CursoFolder.objects.filter(curso=curso, path=path, is_deleted=False).first()
                if not f or f.is_locked:
                    return JsonResponse({"ok": False, "error": "folder_locked_or_missing"}, status=400)

                # собираем содержимое
                has_files_qs = CursoFile.objects.filter(curso=curso, folder_path=path)
                has_sub_qs = CursoFolder.objects.filter(curso=curso, is_deleted=False, path__startswith=(path + "/"))

                files_count = has_files_qs.count()
                sub_count = has_sub_qs.count()

                if mode == "check":
                    if files_count or sub_count:
                        return JsonResponse({
                            "ok": False,
                            "error": "folder_not_empty",
                            "files_count": files_count,
                            "subfolders_count": sub_count,
                        }, status=409)

                    f.is_deleted = True
                    f.save(update_fields=["is_deleted"])
                    return JsonResponse({"ok": True, "deleted": True})

                if mode == "move_root":
                    with transaction.atomic():
                        # 1) переместим файлы из папки и всех подпапок в корень
                        # (если хочешь сохранить структуру — можно иначе, но ты просил “в корень”)
                        CursoFile.objects.filter(
                            curso=curso
                        ).filter(
                            Q(folder_path=path) | Q(folder_path__startswith=(path + "/"))
                        ).update(folder_path="", module_key="")

                        # 2) пометим удалёнными все подпапки и саму папку
                        CursoFolder.objects.filter(curso=curso, is_deleted=False, path__startswith=(path + "/")).update(is_deleted=True)
                        f.is_deleted = True
                        f.save(update_fields=["is_deleted"])

                    return JsonResponse({"ok": True, "deleted": True, "moved_to_root": True})

                if mode == "purge":
                    # ⚠️ аккуратно: если у тебя физические файлы шарятся (copy-private делает ссылку на тот же file),
                    # надо удалять физический файл только когда на него больше нет ссылок.
                    with transaction.atomic():
                        # соберём id файлов, которые надо удалить (в папке и подпапках)
                        qs = CursoFile.objects.filter(curso=curso).filter(
                            Q(folder_path=path) | Q(folder_path__startswith=(path + "/"))
                        )
                        file_ids = list(qs.values_list("id", flat=True))

                        # удаляем записи CursoFile (физику — отдельно)
                        files = list(qs.only("id", "file"))
                        qs.delete()

                        # чистим физические файлы безопасно
                        for cf in files:
                            fname = (cf.file.name or "").strip()
                            if fname and not CursoFile.objects.filter(file=fname).exists():
                                try:
                                    cf.file.storage.delete(fname)
                                except Exception:
                                    pass

                        # пометим папки удалёнными
                        CursoFolder.objects.filter(curso=curso, is_deleted=False, path__startswith=(path + "/")).update(is_deleted=True)
                        f.is_deleted = True
                        f.save(update_fields=["is_deleted"])

                    return JsonResponse({"ok": True, "deleted": True, "purged": True})

                return JsonResponse({"ok": False, "error": "bad_mode"}, status=400)
            
            def _infer_module_from_path(curso, target_path: str) -> str:
                """
                Ищем ближайшую locked папку по пути (самая глубокая).
                Если нашли — module_key = её title (или leaf).
                Если нет — module_key = "".
                """
                tp = (target_path or "").strip().strip("/")
                if not tp:
                    return ""

                parts = tp.split("/")
                prefixes = []
                cur = ""
                for p in parts:
                    cur = f"{cur}/{p}" if cur else p
                    prefixes.append(cur)

                locked = (CursoFolder.objects
                          .filter(curso=curso, is_deleted=False, is_locked=True, path__in=prefixes)
                          .values("path", "title"))

                if not locked:
                    return ""

                # самая глубокая = самая длинная path
                best = sorted(locked, key=lambda x: len(x["path"] or ""), reverse=True)[0]
                title = (best.get("title") or "").strip()
                if title:
                    return title

                # fallback: имя сегмента
                return best["path"].split("/")[-1]

            if action == "file_move":
                fid = request.POST.get("file_id")
                target = _normalize_path(request.POST.get("target_path") or "")

                cf = CursoFile.objects.filter(pk=fid, curso=curso).first()
                if not cf:
                    return JsonResponse({"ok": False, "error": "file_not_found"}, status=404)

                # target="" = корень — всегда ок
                if target:
                    folder = CursoFolder.objects.filter(curso=curso, is_deleted=False, path=target).first()
                    if not folder:
                        return JsonResponse({"ok": False, "error": "folder_not_found"}, status=404)

                    # ✅ ВАЖНО: locked папки модулей НЕ запрещают перенос
                    # if folder.is_locked:  <-- УБРАТЬ БЛОКИРОВКУ

                cf.folder_path = target
                cf.module_key = _infer_module_from_path(curso, target)
                cf.save(update_fields=["folder_path", "module_key"])

                return JsonResponse({
                    "ok": True,
                    "file_id": cf.id,
                    "target_path": target,
                    "module_key": cf.module_key or "",
                })


            if action == "folder_create":
                name = (request.POST.get("folder_name") or "").strip()
                parent = _normalize_path(request.POST.get("folder_parent") or current_path or "")

                try:
                    folder, created = create_folder(
                        curso,
                        parent_path=parent,
                        name=name,
                        user=user,
                        locked=False,
                    )
                except ValueError as e:
                    return JsonResponse({"ok": False, "error": str(e)}, status=400)

                # ✅ соберём node-словарь как в build_node
                node = {
                    "path": folder.path,
                    "name": folder.title or folder.path.split("/")[-1],
                    "is_locked": bool(folder.is_locked),
                    "count": 0,
                    "files": [],
                    "children": [],
                }

                html = render_to_string(
                    "panel/partials/materiales_tree_node.html",
                    {"node": node, "is_teacher": True},
                    request=request,
                )

                if request.headers.get("x-requested-with") == "XMLHttpRequest":
                    return JsonResponse({
                        "ok": True,
                        "path": folder.path,
                        "parent_path": parent,
                        "node_html": html,
                    })

                # fallback non-AJAX
                return redirect(f"{request.path}?tab=materiales&aud={audience}&p={parent}")
                
            if action == "upload_files_ajax":
                folder_path = _normalize_path(request.POST.get("folder_path") or current_path or "")
                files = request.FILES.getlist("files")

                if not files:
                    return JsonResponse({"ok": False, "error": "no_files"}, status=400)

                created_items = []
                for f in files:
                    obj = CursoFile(
                        curso=curso,
                        uploaded_by=user,
                        tipo=CursoFile.TIPO_ALUMNOS,  # или вычисли по audience
                        folder_path=folder_path,
                        title=f.name,
                        size=f.size,
                        ext=(f.name.rsplit(".", 1)[-1].lower() if "." in f.name else ""),
                    )
                    obj.module_key = _infer_module_from_path(curso, folder_path)
                    obj.file = f
                    obj.save()

                    # вернём html элемента файла, чтобы UI обновился без reload
                    html = render_to_string(
                        "panel/partials/materiales_tree_file.html",
                        {"f": obj, "is_teacher": True, "audience": audience},
                        request=request
                    )
                    created_items.append({"id": obj.id, "file_html": html})

                return JsonResponse({"ok": True, "files": created_items})            
                        
                
                
            # upload
            if "upload" in request.POST:
                tipo = request.POST.get("tipo") or CursoFile.TIPO_ALUMNOS
                if tipo not in {t for t, _ in CursoFile.TIPO_CHOICES}:
                    tipo = CursoFile.TIPO_ALUMNOS

                created = 0
                for k, v in request.POST.items():
                    if not k.startswith("url_"):
                        continue

                    suffix = k.split("_", 1)[1]     # "url_1" -> "1"
                    url = (v or "").strip()
                    if not url:
                        continue

                    title = (request.POST.get(f"url_title_{suffix}") or "").strip()
                    mod_key = (request.POST.get(f"url_module_key_{suffix}") or "").strip()

                    try:
                        filename, data = _download_url(url)
                        folder_path = _normalize_path(request.POST.get("folder_path") or current_path or "")
                        obj = CursoFile(
                            curso=curso,
                            uploaded_by=user,
                            tipo=tipo,
                            module_key=mod_key,
                            folder_path=folder_path, 
                            title=title or filename,
                            size=len(data),
                            ext=(filename.rsplit(".", 1)[-1].lower() if "." in filename else ""),
                        )
                        # важно: сначала сохранить, потом file.save (или наоборот — так тоже ок)
                        obj.save()
                        obj.file.save(filename, ContentFile(data), save=True)

                        created += 1
                        # ✅ AJAX: вернуть html файла сразу, чтобы UI обновился без перезагрузки
                        if request.headers.get("x-requested-with") == "XMLHttpRequest":
                            # если mod_key не задан — можно инферить от папки
                            if not obj.module_key:
                                obj.module_key = _infer_module_from_path(curso, folder_path)
                                obj.save(update_fields=["module_key"])

                            file_html = render_to_string(
                                "panel/partials/materiales_tree_file.html",
                                {"f": obj, "is_teacher": True, "audience": audience},
                                request=request
                            )
                            return JsonResponse({
                                "ok": True,
                                "mode": "url_upload",
                                "file_id": obj.id,
                                "folder_path": folder_path,
                                "module_key": obj.module_key or "",
                                "file_html": file_html,
                                "created": created,
                            })
                    except Exception as e:
                        logger.exception("URL upload failed: %s", url)
                        # хочешь — можешь вернуть ошибку вместо тихого редиректа:
                        # return HttpResponseBadRequest(f"URL upload failed: {e}")

                for field_name, f in request.FILES.items():
                    if not field_name.startswith("file_"):
                        continue

                    suffix = field_name.split("_", 1)[1]
                    mod_key = (request.POST.get(f"module_key_{suffix}") or "").strip()
                    title = (request.POST.get(f"title_{suffix}") or "").strip()
                    folder_path = _normalize_path(request.POST.get("folder_path") or current_path or "")
                    obj = CursoFile(
                        curso=curso,
                        uploaded_by=user,
                        tipo=tipo,
                        module_key=mod_key,
                        folder_path=folder_path, 
                        title=title or f.name,
                        file=f,
                        size=f.size,
                        ext=(f.name.rsplit(".", 1)[-1].lower() if "." in f.name else ""),
                    )
                    obj.save()
                    created += 1

            # delete
            # delete
            if "delete_id" in request.POST:
                fid = request.POST.get("delete_id")
                cf = CursoFile.objects.filter(pk=fid, curso=curso).first()
                if cf:
                    file_name = (cf.file.name or "").strip()
                    src_folder = (cf.folder_path or "").strip()

                    cf.delete()

                    if file_name and not CursoFile.objects.filter(file=file_name).exists():
                        try:
                            cf.file.storage.delete(file_name)
                        except Exception:
                            pass

                    # ✅ AJAX response
                    if request.headers.get("x-requested-with") == "XMLHttpRequest":
                        return JsonResponse({
                            "ok": True,
                            "file_id": str(fid),
                            "folder_path": src_folder,
                        })

                if request.headers.get("x-requested-with") == "XMLHttpRequest":
                    return JsonResponse({"ok": False, "error": "file_not_found"}, status=404)

                return redirect(f"{request.path}?tab=materiales&aud={audience}&sort={sort}&p={current_path}")
                
            if action == "upload_folder_bundle":
                files = request.FILES.getlist('files')
                rels  = request.POST.getlist('paths')

                if len(files) != len(rels):
                    return JsonResponse({"ok": False, "error": "bad_payload"}, status=400)

                tipo = CursoFile.TIPO_ALUMNOS  # или из GET/POST

                created_files = 0
                created_folders = set()

                for fobj, rel in zip(files, rels):
                    rel = (rel or "").replace("\\","/").strip("/")
                    if not rel:
                        continue

                    # "Root/Sub/a.pdf" => folder="Root/Sub"
                    folder = rel.rsplit("/", 1)[0] if "/" in rel else ""
                    folder = normalize_path(folder)

                    if folder and folder not in created_folders:
                        # создаём цепочку родителей тоже
                        cur = ""
                        for seg in folder.split("/"):
                            cur = f"{cur}/{seg}" if cur else seg
                            create_folder(curso, parent_path=parent_path(cur), name=leaf_name(cur), user=user, locked=False)
                        created_folders.add(folder)

                    obj = CursoFile(
                        curso=curso,
                        uploaded_by=user,
                        tipo=tipo,
                        folder_path=folder,
                        module_key=_infer_module_from_path(curso, folder),
                        title=fobj.name,
                        size=fobj.size,
                        ext=(fobj.name.rsplit(".",1)[-1].lower() if "." in fobj.name else ""),
                    )
                    obj.save()
                    obj.file.save(fobj.name, fobj, save=True)
                    created_files += 1

                return JsonResponse({"ok": True, "files_created": created_files, "folders": len(created_folders)})

            # share/unshare
            if "share_id" in request.POST:
                fid = request.POST.get("share_id")
                cf = CursoFile.objects.filter(pk=fid, curso=curso).first()
                if cf:
                    cf.share_alumnos = not cf.share_alumnos
                    cf.save(update_fields=["share_alumnos"])
                return redirect(f"{request.path}?tab=materiales&aud={audience}&sort={sort}")

            # copy to privados
            if "copy_id" in request.POST:
                fid = request.POST.get("copy_id")
                src = CursoFile.objects.filter(pk=fid, curso=curso).first()
                if src:
                    new = CursoFile(
                        curso=curso,
                        uploaded_by=user,
                        tipo=CursoFile.TIPO_PRIVADO,
                        module_key=src.module_key,
                        title=src.title,
                        file=src.file,
                        size=src.size,
                        ext=src.ext,
                        share_alumnos=False,
                    )
                    new.save()
                return redirect(f"{request.path}?tab=materiales&aud={audience}&sort={sort}")

        # ===== counts for filters BEFORE module filter =====
        from collections import Counter
        files_for_counts = list(files_qs)
        counts_by_mod = Counter((f.module_key or "").strip() for f in files_for_counts)

        total_files = len(files_for_counts)
        none_count = counts_by_mod.get("", 0)

        # ===== mod_groups MF → UF =====
        mod_groups = []
        for mod in tmp:
            if mod.get("kind") != "MF":
                continue
            mf_name = (mod.get("name") or "").strip()
            if not mf_name:
                continue

            items = [{"name": mf_name, "count": int(counts_by_mod.get(mf_name, 0))}]
            for c_idx in (mod.get("children") or []):
                uf_name = (tmp[c_idx].get("name") or "").strip()
                if not uf_name:
                    continue
                items.append({"name": uf_name, "count": int(counts_by_mod.get(uf_name, 0))})

            mod_groups.append({"mf": mf_name, "items": items, "total": sum(x["count"] for x in items)})

        # ===== apply module filter to actual list =====
        if selected_mod == "none":
            files_qs = files_qs.filter(module_key="")
        elif selected_mod:
            files_qs = files_qs.filter(module_key=selected_mod)

        # ===== base ordering inside module =====
        # created_at должен быть в модели CursoFile (у тебя он уже используется выше)
        if sort == "new":
            files_qs = files_qs.order_by("-created_at", "-id")
        elif sort == "old":
            files_qs = files_qs.order_by("created_at", "id")
        elif sort == "az":
            files_qs = files_qs.order_by("title", "id")
        elif sort == "za":
            files_qs = files_qs.order_by("-title", "-id")

        files = list(files_qs)


        # downloaded marker for students
        if not is_teacher:
            downloaded_ids = set(
                MaterialDownload.objects.filter(alumno=user, file__in=files)
                .values_list("file_id", flat=True)
            )
            for f in files:
                f.is_downloaded = (f.id in downloaded_ids)

        # sorting by module order (from INFO list)
        mod_order = {(m["name"] or "").strip(): idx for idx, m in enumerate(mods, start=1)}

        def _module_group_key(f):
            key = (f.module_key or "").strip()
            order = 0 if not key else mod_order.get(key, 9999)
            return (order, key.lower())

        # Python sort стабильный => порядок внутри одной группы сохранится как пришёл из DB (created_at sort)
        files.sort(key=_module_group_key)


        # fmt_size
        for f in files:
            f.fmt_size = _fmt_size(getattr(f, "size", 0))

        mods_with_counts = []
        for m in mods:
            name = (m["name"] or "").strip()
            if name:
                mods_with_counts.append({"name": name, "count": int(counts_by_mod.get(name, 0))})

        context.update({
            "files": files,
            "audience": audience,
            "selected_mod": selected_mod,
            "total_files": total_files,
            "none_count": none_count,
            "mods_with_counts": mods_with_counts,
            "mod_groups": mod_groups,
        })

        # ===== Receipts context for student =====
        if not is_teacher:
            qs = MaterialReceipt.objects.filter(curso=curso, alumno=user)
            receipt_keys = set(qs.values_list("item_key", flat=True))
            receipt_locked_keys = set(receipt_keys)
            keys_all = {it.key for it in physical_items_enabled}

            context.update({
                "receipt_keys": receipt_keys,
                "receipt_locked_keys": receipt_locked_keys,
                "receipt_all_done": (len(keys_all) > 0 and receipt_keys.issuperset(keys_all)),
            })

    # ───────────── ВКЛАДКА "TAREAS" ─────────────
    elif active_tab == "tareas":

        def _fmt_size(num):
            try:
                n = float(num or 0)
            except (TypeError, ValueError):
                n = 0.0
            for unit in ("B", "KB", "MB", "GB", "TB"):
                if n < 1024.0 or unit == "TB":
                    if unit == "B":
                        return f"{int(n)} B"
                    return f"{n:.1f} {unit}"
                n /= 1024.0
            return f"{n} B"
        # choices для селекта
        module_choices = curso_modules_choices(curso)  # [(key,label), ...]
        context["module_choices"] = module_choices

        # выбранный модуль из GET (Todos / конкретный / none)
        selected_mod = (request.GET.get("mod") or "").strip()  # "" = Todos
        context["selected_mod"] = selected_mod

        # валидные ключи
        valid_mod_keys = {k for k, _ in module_choices}

        # удобные флаги для UI
        has_none = CourseTask.objects.filter(curso=curso, module_key="").exists()
        context["has_none_mod"] = has_none


        # ── TEACHER: создать задачу
        if request.method == "POST" and is_teacher and request.POST.get("action") == "task_create":
            title = (request.POST.get("title") or "").strip()
            description = (request.POST.get("description") or "").strip()
            f = request.FILES.get("task_file")
            module_key = (request.POST.get("module_key") or "").strip()

            # NEW:
            is_final_exam = (request.POST.get("is_final_exam") == "on")
            conv_raw = (request.POST.get("convocatoria") or "").strip()
            convocatoria = None
            if is_final_exam and conv_raw in {"1", "2"}:
                convocatoria = int(conv_raw)

            valid = {k for k, _ in curso_modules_choices(curso)}
            if module_key not in valid:
                return HttpResponseBadRequest("Módulo inválido")

            if title:
                t = CourseTask.objects.create(
                    curso=curso,
                    created_by=user,
                    title=title,
                    description=description,
                    module_key=module_key,
                    module_label=curso_module_label(curso, module_key),
                    is_published=True,
                    is_closed=False,

                    # NEW:
                    is_final_exam=is_final_exam,
                    convocatoria=convocatoria,
                )
                if f:
                    t.file = f
                    t.file_name = f.name
                    t.file_size = f.size
                    t.ext = (f.name.rsplit(".", 1)[-1].lower() if "." in f.name else "")
                    t.save(update_fields=["file", "file_name", "file_size", "ext"])

            return redirect(f"{request.path}?tab=tareas")


        # ── TEACHER: закрыть/открыть задачу
        if request.method == "POST" and is_teacher and request.POST.get("action") == "task_toggle_close":
            task_id = request.POST.get("task_id")
            t = CourseTask.objects.filter(pk=task_id, curso=curso).first()
            if t:
                t.is_closed = not bool(t.is_closed)
                t.save(update_fields=["is_closed"])
            return redirect(f"{request.path}?tab=tareas")
            
        # ── TEACHER: удалить задачу целиком
        if request.method == "POST" and is_teacher and request.POST.get("action") == "task_delete":
            task_id = request.POST.get("task_id")
            t = CourseTask.objects.filter(pk=task_id, curso=curso).first()
            if t:
                # удаляем файл задания
                if t.file:
                    t.file.delete(save=False)

                # удаляем файлы сдач + записи
                subs = TaskSubmission.objects.filter(task=t)
                for s in subs:
                    if s.file:
                        s.file.delete(save=False)
                subs.delete()

                t.delete()

            return redirect(f"{request.path}?tab=tareas")


        # ── STUDENT: отправить работу (можно несколько раз)
        # ── STUDENT: отправить работу (много попыток ДО оценки)
        if request.method == "POST" and (not is_teacher) and request.POST.get("action") == "task_submit":
            task_id = request.POST.get("task_id")
            f = request.FILES.get("answer_file")
            comment = (request.POST.get("comment") or "").strip()

            task = CourseTask.objects.filter(pk=task_id, curso=curso, is_published=True).first()

            if task and (not task.is_closed) and f:
                sub = TaskSubmission.objects.filter(task=task, alumno=user).first()

                # ✅ если уже оценено — блокируем любые новые попытки
                if sub and sub.status == TaskSubmission.STATUS_GRADED:
                    return redirect(f"{request.path}?tab=tareas")

                # иначе: обновляем существующую сдачу или создаём новую
                if sub:
                    if sub.file:
                        try:
                            sub.file.delete(save=False)
                        except Exception:
                            pass

                    sub.file = f
                    sub.file_name = f.name
                    sub.file_size = f.size
                    sub.ext = (f.name.rsplit(".", 1)[-1].lower() if "." in f.name else "")
                    sub.comment = comment
                    sub.status = TaskSubmission.STATUS_SUBMITTED
                    sub.submitted_at = timezone.now()
                    sub.save()
                else:
                    TaskSubmission.objects.create(
                        task=task,
                        alumno=user,
                        file=f,
                        file_name=f.name,
                        file_size=f.size,
                        ext=(f.name.rsplit(".", 1)[-1].lower() if "." in f.name else ""),
                        comment=comment,
                        status=TaskSubmission.STATUS_SUBMITTED,
                        submitted_at=timezone.now(),
                    )

            return redirect(f"{request.path}?tab=tareas")



        # ── TEACHER: поставить оценку (ТОЛЬКО 1 РАЗ, потом блок)
        if request.method == "POST" and is_teacher and request.POST.get("action") == "task_grade":
            sub_id = request.POST.get("sub_id")
            grade_raw = (request.POST.get("grade") or "").replace(",", ".").strip()
            feedback = (request.POST.get("teacher_feedback") or "").strip()

            sub = TaskSubmission.objects.select_related("task").filter(
                pk=sub_id, task__curso=curso
            ).first()

            if sub and sub.status != TaskSubmission.STATUS_GRADED:  # ✅ не даём менять после оценки
                sub.teacher_feedback = feedback
                sub.graded_by = user
                sub.graded_at = timezone.now()
                sub.status = TaskSubmission.STATUS_GRADED

                try:
                    sub.grade = float(grade_raw) if grade_raw else None
                except ValueError:
                    sub.grade = None

                sub.save()

            return redirect(f"{request.path}?tab=tareas")

        # ── queries + контекст
        if is_teacher:
            tasks_qs = CourseTask.objects.filter(curso=curso)

            if selected_mod == "none":
                tasks_qs = tasks_qs.filter(module_key="")
            elif selected_mod:
                # защита от мусора в url
                if selected_mod in valid_mod_keys:
                    tasks_qs = tasks_qs.filter(module_key=selected_mod)
                else:
                    # если пришёл невалидный ключ — считаем как Todos
                    selected_mod = ""
                    context["selected_mod"] = ""

            tasks = list(
                tasks_qs.order_by("-created_at")
                .prefetch_related(
                    Prefetch(
                        "submissions",
                        queryset=TaskSubmission.objects.select_related("alumno").order_by("status", "-submitted_at")
                    )
                )
            )


            for t in tasks:
                t.fmt_size = _fmt_size(getattr(t, "file_size", 0))
                for s in t.submissions.all():
                    s.fmt_size = _fmt_size(getattr(s, "file_size", 0))

            context["tasks"] = tasks

        else:
            # ✅ STUDENT видит опубликованные задачи
            tasks_qs = CourseTask.objects.filter(curso=curso, is_published=True)

            if selected_mod == "none":
                tasks_qs = tasks_qs.filter(module_key="")
            elif selected_mod:
                if selected_mod in valid_mod_keys:
                    tasks_qs = tasks_qs.filter(module_key=selected_mod)
                else:
                    selected_mod = ""
                    context["selected_mod"] = ""

            tasks = list(tasks_qs.order_by("-created_at"))


            # ✅ его сдачи (нужно для t.my_sub в шаблоне)
            subs_qs = (
                TaskSubmission.objects
                .filter(alumno=user, task__in=tasks)
                .select_related("task")
                .order_by("task_id", "-submitted_at", "-id")
            )

            my_subs = {}
            for s in subs_qs:
                if s.task_id not in my_subs:  # ✅ берём только самую свежую по каждой задаче
                    my_subs[s.task_id] = s


            for t in tasks:
                t.fmt_size = _fmt_size(getattr(t, "file_size", 0))
                t.my_sub = my_subs.get(t.id)
                if t.my_sub:
                    t.my_sub.fmt_size = _fmt_size(getattr(t.my_sub, "file_size", 0))

            context["tasks"] = tasks


    # ───────────── ВКЛАДКА "CALENDARIO" ─────────────
    elif active_tab == "calendario":
        horarios = (
            Horario.objects
            .filter(curso=curso)
            .order_by("dia", "hora_inicio")
        )

        PHYSICAL_KEYS = list(
            CursoPhysicalItem.objects.filter(curso=curso, is_enabled=True)
            .values_list("key", flat=True)
        )



        # какие файлы видимы ученикам в этом курсе
        visible_files_qs = CursoFile.objects.filter(
            curso=curso
        ).filter(
            models.Q(tipo=CursoFile.TIPO_ALUMNOS) |
            (models.Q(share_alumnos=True) & models.Q(tipo__in=[CursoFile.TIPO_DOCENTES, CursoFile.TIPO_PRIVADO]))
        )
        visible_file_ids = list(visible_files_qs.values_list("id", flat=True))
        total_visible = len(visible_file_ids)

        # скачивания: сколько файлов скачал каждый alumno (уникально по file+alumno)
        dl_map = dict(
            MaterialDownload.objects.filter(file_id__in=visible_file_ids)
            .values("alumno_id")
            .annotate(c=Count("id"))
            .values_list("alumno_id", "c")
        )

        # физматериалы: сколько галочек из 3 поставил
        rc_map = dict(
            MaterialReceipt.objects.filter(curso=curso, item_key__in=PHYSICAL_KEYS)
            .values("alumno_id")
            .annotate(c=Count("id"))
            .values_list("alumno_id", "c")
        )

        # теперь добавляем в students
        for s in students:
            s.total_files = total_visible
            s.downloaded_files = int(dl_map.get(s.id, 0))
            s.receipts_done = (int(rc_map.get(s.id, 0)))




        # данные для JS-календаря
        horarios_data = []
        for h in horarios:
            horarios_data.append({
                "dia": h.dia.isoformat(),                  # '2025-12-22'
                "hora_inicio": h.hora_inicio.strftime("%H:%M"),
                "hora_fin": h.hora_fin.strftime("%H:%M"),
                "aula": h.aula,
                "modulo": h.modulo,
                "tipo": h.tipo,
                "grupo": h.grupo,
            })

        context["horarios"] = horarios
        context["horarios_data"] = horarios_data
        
        # ───────────── ВКЛАДКА "ALUMNOS" ─────────────
    elif active_tab == "alumnos":
        # только преподаватель видит эту вкладку (guard стоит и в меню)
        if not is_teacher:
            context["students"] = []
        else:
            # кэш последних сгенерированных паролей в сессии
            last_passes = request.session.get("panel_last_passes", {})

            # обработка POST: reset / remove
            if request.method == "POST":
                # удалить ученика с курса (и при необходимости удалить user целиком)
                if "remove_id" in request.POST:
                    uid = request.POST.get("remove_id")
                    try:
                        uid_int = int(uid)
                    except (TypeError, ValueError):
                        uid_int = None

                    if uid_int:
                        # 1) снимаем с курса
                        Enrol.objects.filter(
                            user_id=uid_int,
                            codigo=curso.codigo,
                            role="student",
                        ).delete()

                        # 2) если у юзера больше нет enrol-ов вообще — удаляем полностью
                        other_enrols = Enrol.objects.filter(user_id=uid_int).exists()
                        if not other_enrols:
                            # чистим связанные данные (на всякий случай, даже если FK CASCADE)
                            MaterialDownload.objects.filter(alumno_id=uid_int).delete()
                            MaterialReceipt.objects.filter(alumno_id=uid_int).delete()
                            TaskSubmission.objects.filter(alumno_id=uid_int).delete()
                            AttendanceSession.objects.filter(user_id=uid_int).delete()

                            # профиль (если есть)
                            if UserProfile is not None:
                                UserProfile.objects.filter(user_id=uid_int).delete()

                            # сам пользователь
                            User.objects.filter(id=uid_int).delete()

                        # уберём и из кэша паролей
                        last_passes.pop(str(uid_int), None)

                    request.session["panel_last_passes"] = last_passes
                    return redirect(f"{request.path}?tab=alumnos")


                # сбросить пароль ученику
                if "reset_id" in request.POST:
                    uid = request.POST.get("reset_id")
                    try:
                        uid_int = int(uid)
                    except (TypeError, ValueError):
                        uid_int = None

                    if uid_int:
                        try:
                            u = User.objects.get(pk=uid_int)
                        except User.DoesNotExist:
                            pass
                        else:
                            # генерим новый простой пароль (можешь поменять длину/алфавит)
                            new_pwd = f"{random.randint(0, 9999):04d}"
                            u.set_password(new_pwd)
                            u.save(update_fields=["password"])
                            last_passes[str(uid_int)] = new_pwd

                    request.session["panel_last_passes"] = last_passes
                    return redirect(f"{request.path}?tab=alumnos")

            # GET: собираем список студентов
            enrols = (
                Enrol.objects
                .filter(codigo=curso.codigo, role="student")
                .select_related("user")
                .order_by("created_at")
            )

            students = []

            for idx, enr in enumerate(enrols, start=1):
                u = enr.user

                # пробуем достать профиль
                profile = getattr(u, "profile", None)  # если related_name="profile"
                apellidos = " ".join([
                    (getattr(profile, "last_name1", "") or "").strip(),
                    (getattr(profile, "last_name2", "") or "").strip(),
                ]).strip()
                nombre = (getattr(profile, "first_name", "") or u.first_name or "").strip()

                display = f"{apellidos}, {nombre}".strip(", ").strip() or (u.get_full_name() or u.email or u.username or "")


                if profile and getattr(profile, "display_name", ""):
                    display_name = profile.display_name.strip()
                else:
                    display_name = (
                        (u.get_full_name() or "").strip()
                        or (u.username or "") 
                        or (u.email or "")
                        or f"ID {u.pk}"
                    )

                students.append({
                    "idx": idx,
                    "id": u.pk,
                    "email": u.email,
                    "name": display_name,
                    "enrol_created_at": enr.created_at,
                    "last_pass": last_passes.get(str(u.pk)),
                })


            context.update({
                "students": students,
            })

            PHYSICAL_KEYS = list(
                CursoPhysicalItem.objects.filter(curso=curso, is_enabled=True)
                .values_list("key", flat=True)
            )

            visible_files_qs = CursoFile.objects.filter(curso=curso).filter(
                Q(tipo=CursoFile.TIPO_ALUMNOS) |
                (Q(share_alumnos=True) & Q(tipo__in=[CursoFile.TIPO_DOCENTES, CursoFile.TIPO_PRIVADO]))
            )
            visible_ids = list(visible_files_qs.values_list("id", flat=True))
            total_visible = len(visible_ids)

            dl_map = dict(
                MaterialDownload.objects.filter(file_id__in=visible_ids)
                .values("alumno_id").annotate(c=Count("id"))
                .values_list("alumno_id", "c")
            )

            rc_map = dict(
                MaterialReceipt.objects.filter(curso=curso, item_key__in=PHYSICAL_KEYS)
                .values("alumno_id").annotate(c=Count("id"))
                .values_list("alumno_id", "c")
            )

            for s in students:
                sid = s["id"]
                s["total_files"] = total_visible
                s["downloaded_files"] = int(dl_map.get(sid, 0))
                s["receipts_count"] = int(rc_map.get(sid, 0))
                
            # --- AVG grade per alumno (only graded submissions) ---
            graded = (
                TaskSubmission.objects
                .filter(task__curso=curso, status=TaskSubmission.STATUS_GRADED)
                .values("alumno_id")
                .annotate(
                    avg=Avg("grade"),
                    n=Count("id", filter=Q(grade__isnull=False)),
                )
            )

            grade_map = {row["alumno_id"]: row for row in graded}

            for s in students:
                row = grade_map.get(s["id"])
                if row and row["avg"] is not None:
                    # округление как хочешь: 1 знак или 2
                    s["avg_grade"] = round(float(row["avg"]), 1)
                    s["graded_tasks"] = int(row["n"] or 0)
                else:
                    s["avg_grade"] = None
                    s["graded_tasks"] = 0


    return render(request, "panel/course_panel.html", context)

@login_required
@require_profile_complete
def alumno_home(request):
    user = request.user
    is_teacher = Enrol.objects.filter(user=user, role="teacher").exists()

    # --- курсы через Enrol ---
    if is_teacher:
        # все курсы, где он назначен как teacher
        codigos = (
            Enrol.objects
            .filter(user=user, role="teacher")
            .values_list("codigo", flat=True)
            .distinct()
        )
    else:
        # все курсы, где он ученик (любая роль, кроме teacher)
        codigos = (
            Enrol.objects
            .filter(user=user)
            .exclude(role="teacher")
            .values_list("codigo", flat=True)
            .distinct()
        )

    cursos = list(
        Curso.objects.filter(codigo__in=codigos).order_by("codigo")
    )

    # если ровно один курс — сразу кидаем в панель этого курса
    if len(cursos) == 1:
        codigo = (cursos[0].codigo or "").strip()
        return redirect("panel:course", codigo=codigo)


    # иначе показываем список курсов
    return render(request, "panel/alumno_home.html", {
        "cursos": cursos,
        "is_teacher": is_teacher,
    })
    
from django.http import FileResponse, Http404
def get_client_ip(request):
    xff = request.META.get("HTTP_X_FORWARDED_FOR")
    if xff:
        return xff.split(",")[0].strip()
    return request.META.get("REMOTE_ADDR")

def user_can_access_course(user, curso: Curso) -> bool:
    # teacher
    if Enrol.objects.filter(user=user, codigo=curso.codigo, role="teacher").exists():
        return True
    # student
    if Enrol.objects.filter(user=user, codigo=curso.codigo).exclude(role="teacher").exists():
        return True
    return False


from django.views.decorators.clickjacking import xframe_options_sameorigin
import mimetypes

@login_required
@xframe_options_sameorigin
def task_download(request, codigo: str, task_id: int):
    curso = get_object_or_404(Curso, codigo=codigo)
    if not user_can_access_course(request.user, curso):
        raise Http404

    task = CourseTask.objects.filter(pk=task_id, curso=curso, is_published=True).first()
    if not task or not task.file:
        raise Http404

    inline = request.GET.get("inline") == "1"
    ctype, _ = mimetypes.guess_type(task.filename or "")
    ctype = ctype or "application/octet-stream"

    resp = FileResponse(task.file.open("rb"), content_type=ctype)
    filename = task.filename or "task"

    resp["Content-Disposition"] = (
        f'inline; filename="{filename}"' if inline
        else f'attachment; filename="{filename}"'
    )
    return resp


from django.views.decorators.clickjacking import xframe_options_sameorigin
import mimetypes
from django.http import FileResponse, Http404

@login_required
@xframe_options_sameorigin
def submission_download(request, codigo: str, task_id: int, sub_id: int):
    curso = get_object_or_404(Curso, codigo=codigo)
    if not user_can_access_course(request.user, curso):
        raise Http404

    is_teacher = Enrol.objects.filter(user=request.user, codigo=curso.codigo, role="teacher").exists()

    sub = (TaskSubmission.objects
           .select_related("task")
           .filter(pk=sub_id, task_id=task_id, task__curso=curso)
           .first())
    if not sub or not sub.file:
        raise Http404

    # ученик — только свой файл
    if (not is_teacher) and (sub.alumno_id != request.user.id):
        raise Http404

    inline = request.GET.get("inline") == "1"
    ctype, _ = mimetypes.guess_type(sub.filename or "")
    ctype = ctype or "application/octet-stream"

    resp = FileResponse(sub.file.open("rb"), content_type=ctype)
    filename = sub.filename or "submission"

    resp["Content-Disposition"] = (
        f'inline; filename="{filename}"' if inline
        else f'attachment; filename="{filename}"'
    )
    return resp


# api/views_me.py (примерно)
from rest_framework.decorators import api_view, permission_classes
from rest_framework.permissions import IsAuthenticated
from rest_framework.response import Response
from api.models import UserProfile

@api_view(["GET", "POST"])
@permission_classes([IsAuthenticated])
def me_profile(request):
    p, _ = UserProfile.objects.get_or_create(user=request.user)

    if request.method == "GET":
        return Response({
            "profile": {
                "first_name": p.first_name,
                "last_name1": p.last_name1,
                "last_name2": p.last_name2,
                "display_name": p.display_name,
                "bio": p.bio,
                "is_teacher": p.is_teacher,
                "is_complete": p.is_complete(),
            }
        })

    data = request.data or {}
    p.first_name   = (data.get("first_name") or "").strip()
    p.last_name1   = (data.get("last_name1") or "").strip()
    p.last_name2   = (data.get("last_name2") or "").strip()
    p.display_name = (data.get("display_name") or "").strip()
    p.bio          = (data.get("bio") or "").strip()

    # ✅ если display пустой — сгенерим из имени/фамилий
    if not p.display_name:
        gen = p.build_display_name()
        if gen:
            p.display_name = gen

    p.save()
    return Response({"ok": True, "profile": {
        "first_name": p.first_name,
        "last_name1": p.last_name1,
        "last_name2": p.last_name2,
        "display_name": p.display_name,
        "bio": p.bio,
        "is_complete": p.is_complete(),
    }})


@csrf_exempt
@require_http_methods(["POST"])
def material_share_link_create(request, file_id: int):
    """
    POST /panel/materiales/<file_id>/share-link
    -> {"ok": true, "url": "...", "expires_at": "..."}
    """
    user = request.user
    if not user.is_authenticated:
        return JsonResponse({"ok": False, "error": "auth_required"}, status=401)

    # твоя логика: teacher/admin
    curso = cf.curso
    is_teacher = Enrol.objects.filter(user=user, codigo=curso.codigo, role="teacher").exists() or user.is_staff
    if not is_teacher:
        return JsonResponse({"ok": False, "error": "forbidden"}, status=403)

    cf = CursoFile.objects.select_related("curso").filter(pk=file_id).first()
    if not cf:
        return JsonResponse({"ok": False, "error": "not_found"}, status=404)

    # ⚠️ Политика: какие файлы можно отдавать публично?
    # Я бы разрешил только то, что видимо ученикам:
    # - alumnos
    # - или share_alumnos=True (docentes/privado)
    allowed_public = (cf.tipo == "alumnos") or (cf.share_alumnos and cf.tipo in ("docentes", "privado"))
    if not allowed_public:
        return JsonResponse({"ok": False, "error": "not_publicable"}, status=400)

    # срок жизни (например, 14 дней). Можно сделать настройкой.
    ttl_days = 14
    expires_at = timezone.now() + timedelta(days=ttl_days)

    # создаём новый токен (или можно "переиспользовать" активный)
    link = PublicShareLink.objects.create(
        token=PublicShareLink.new_token(),
        file=cf,
        created_by=user,
        expires_at=expires_at,
        is_active=True,
    )

    # абсолютная ссылка
    url = request.build_absolute_uri(f"/alumno/s/{link.token}/")


    return JsonResponse({
        "ok": True,
        "url": url,
        "expires_at": link.expires_at.isoformat() if link.expires_at else None
    })

from django.shortcuts import get_object_or_404
from django.views.decorators.http import require_GET

@require_GET
def public_share_download(request, token: str):
    """
    GET /s/<token>
    -> отдаёт файл как download, без логина
    """
    link = PublicShareLink.objects.select_related("file").filter(token=token).first()
    if not link or not link.is_valid():
        raise Http404()

    cf = link.file

    # Доп. защита: если после создания ссылки файл стал "не видим ученикам" — блокируем.
    allowed_public = (cf.tipo == "alumnos") or (cf.share_alumnos and cf.tipo in ("docentes", "privado"))
    if not allowed_public:
        raise Http404()

    # Здесь можно логировать скачивания анонимных (по желанию)
    # MaterialDownload.objects.create(... user=None, ...)

    # Отдаём файл
    f = cf.file
    if not f:
        raise Http404()

    resp = HttpResponse(f.open("rb").read(), content_type="application/octet-stream")
    filename = cf.filename or "material"
    resp["Content-Disposition"] = f'attachment; filename="{filename}"'
    return resp


from django.http import FileResponse, Http404
from django.views.decorators.clickjacking import xframe_options_sameorigin
import mimetypes

from django.db import IntegrityError

@login_required
@xframe_options_sameorigin
def material_download(request, file_id: int):
    f = CursoFile.objects.select_related("curso").filter(id=file_id).first()
    if not f:
        raise Http404

    curso = f.curso
    if not user_can_access_course(request.user, curso):
        raise Http404

    is_teacher = Enrol.objects.filter(user=request.user, codigo=curso.codigo, role="teacher").exists()
    if not f.can_see(request.user, is_teacher=is_teacher):
        raise Http404

    inline = request.GET.get("inline") == "1"

    # content-type + отдача
    ctype, _ = mimetypes.guess_type(f.filename or "")
    ctype = ctype or "application/octet-stream"

    resp = FileResponse(f.file.open("rb"), content_type=ctype)
    filename = f.filename or "file"
    resp["Content-Disposition"] = (
        f'inline; filename="{filename}"' if inline
        else f'attachment; filename="{filename}"'
    )
        # ✅ ЛОГ СКАЧИВАНИЯ: только для ученика, только при attachment (не inline preview)
    if (not is_teacher) and (not inline):
        try:
            MaterialDownload.objects.get_or_create(
                file=f,
                alumno=request.user,
                defaults={
                    "ip": get_client_ip(request),
                    "user_agent": (request.META.get("HTTP_USER_AGENT") or "")[:255],
                }
            )
        except Exception:
            # никогда не ломаем выдачу файла из-за логов
            pass
    return resp


# ============================================================
# DOCX: Acta entrega material físico (simple, no duplicates)
# ============================================================
from pathlib import Path
from datetime import date as _date

from django.conf import settings

from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def _clear_container(container):
    # удаляем таблицы и параграфы из header/footer
    for t in list(container.tables):
        t._tbl.getparent().remove(t._tbl)
    for p in list(container.paragraphs):
        p._p.getparent().remove(p._p)

def _set_table_width_twips(table, tw: int):
    """
    Жёстко задаём общую ширину таблицы (tblW) в twips.
    """
    tblPr = table._tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        table._tbl.insert(0, tblPr)

    # remove existing tblW
    for el in tblPr.findall(qn("w:tblW")):
        tblPr.remove(el)

    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:type"), "dxa")
    tblW.set(qn("w:w"), str(int(tw)))
    tblPr.append(tblW)


def _apply_col_widths_to_all_rows(table, widths_twips: list[int]):
    """
    Word любит "забывать" ширины колонок на добавленных строках.
    Поэтому ставим width каждой ячейке в каждой строке.
    """
    # columns
    for i, w in enumerate(widths_twips):
        try:
            table.columns[i].width = Twips(int(w))
        except Exception:
            pass

    # every row / cell
    for row in table.rows:
        for i, w in enumerate(widths_twips):
            row.cells[i].width = Twips(int(w))

def _set_table_layout_fixed(table):
    tblPr = table._tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        table._tbl.insert(0, tblPr)

    # remove existing
    for el in tblPr.findall(qn("w:tblLayout")):
        tblPr.remove(el)

    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)


def _set_table_indent_zero(table):
    tblPr = table._tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        table._tbl.insert(0, tblPr)

    for el in tblPr.findall(qn("w:tblInd")):
        tblPr.remove(el)

    tblInd = OxmlElement("w:tblInd")
    tblInd.set(qn("w:w"), "0")
    tblInd.set(qn("w:type"), "dxa")
    tblPr.append(tblInd)


def inject_headers_footers_original(docx_path: Path):
    """
    ОРИГИНАЛЬНАЯ шапка как у тебя:
    - Header: слева Lanbide (2.5cm), справа Euskadi (9.5cm)
    - Footer: справа EU (3.2cm)
    """
    LOGOS = getattr(settings, "MEATZE_DOCX_LOGOS", {}) or {}
    doc = Document(str(docx_path))
    section = doc.sections[0]

    # поля секции как у тебя
    section.left_margin  = Cm(1.0)
    section.right_margin = Cm(1.0)
    section.top_margin   = Cm(1.2)
    section.bottom_margin= Cm(1.2)

    usable_width = section.page_width - section.left_margin - section.right_margin
    tw = _to_twips(usable_width) 

    # ===== HEADER =====
    header = section.header
    _clear_container(header)

    h_table = header.add_table(rows=1, cols=2, width=usable_width)
    h_table.autofit = False
    h_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_indent_zero(h_table)
    _set_table_layout_fixed(h_table)

    # ✅ правильные ширины: 50/50
    h_table.columns[0].width = int(tw * 0.5)
    h_table.columns[1].width = int(tw * 0.5)

    cell_left, cell_right = h_table.rows[0].cells

    # lanbide left
    lanbide = (LOGOS.get("lanbide") or "").strip()
    if lanbide and Path(lanbide).exists():
        p_l = cell_left.paragraphs[0]
        p_l.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_l.paragraph_format.space_before = Pt(0)
        p_l.paragraph_format.space_after  = Pt(0)
        p_l.add_run().add_picture(lanbide, width=Cm(2.5))

    # euskadi right
    euskadi = (LOGOS.get("euskadi") or "").strip()
    if euskadi and Path(euskadi).exists():
        p_r = cell_right.paragraphs[0]
        p_r.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_r.paragraph_format.space_before = Pt(0)
        p_r.paragraph_format.space_after  = Pt(0)
        p_r.add_run().add_picture(euskadi, width=Cm(9.5))

    # ===== FOOTER =====
    footer = section.footer
    _clear_container(footer)

    section.footer_distance = Cm(1.2)

    f_table = footer.add_table(rows=1, cols=1, width=usable_width)
    f_table.autofit = False
    _set_table_indent_zero(f_table)
    _set_table_layout_fixed(f_table)

    eu = (LOGOS.get("eu") or "").strip()
    if eu and Path(eu).exists():
        p = f_table.rows[0].cells[0].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        p.add_run().add_picture(eu, width=Cm(3.2))

    doc.save(str(docx_path))
from pathlib import Path
from datetime import date as _date

from docx import Document
from docx.shared import Cm, Pt, Twips  # ✅ FIX: добавили Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.table import WD_ALIGN_VERTICAL  # если у тебя не импортирован

# ------------------------------------------------------------------

def _to_twips(x) -> int:
    """
    python-docx Length -> twips.

    ВАЖНО:
    - В некоторых версиях python-docx (и в некоторых сборках) section.page_width/margins
      могут приходить как int (EMU). Тогда надо переводить EMU -> twips.
    """
    if x is None:
        return 0

    # Length objects
    if hasattr(x, "twips"):
        return int(x.twips)
    if hasattr(x, "emu"):
        return int(x.emu // 635)  # 1 twip = 635 EMU

    # Raw numbers: почти всегда это EMU (если число большое)
    if isinstance(x, (int, float)):
        n = int(x)
        # эвристика: twips для A4 usable width ~ 9000-11000,
        # EMU обычно миллионы
        if n > 200000:   # значит это EMU
            return int(n // 635)
        return n  # если вдруг реально twips (малое число)

    return int(x)


def _zwsp_wrap_email(email: str) -> str:
    s = (email or "").strip()
    if not s:
        return ""
    zwsp = "\u200b"
    return s.replace("@", "@"+zwsp).replace(".", "."+zwsp)

def _apply_cell_style(cell, *, font_pt=10, bold=False, align=None):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    for r in p.runs:
        r.font.size = Pt(font_pt)
        if bold:
            r.bold = True

# ------------------------------------------------------------------

def build_physical_report_docx(curso, item_label: str, alumnos_rows, fecha: _date, out_path: Path):
    doc = Document()
    section = doc.sections[0]

    section.left_margin = Cm(1.0)
    section.right_margin = Cm(1.0)
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)

    usable_width = section.page_width - section.left_margin - section.right_margin
    tw = _to_twips(usable_width)  # ✅ FIX: twips всегда int

    # ========= Карточка курса (2 колонки) =========
    t = doc.add_table(rows=4, cols=2)
    t.style = "Table Grid"
    t.autofit = False
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_indent_zero(t)
    _set_table_layout_fixed(t)

    # ✅ FIX: ширины ставим через Twips(...)
    t.columns[0].width = Twips(int(tw * 0.26))
    t.columns[1].width = Twips(int(tw * 0.74))

    rows = [
        ("NOMBRE DEL CURSO", (getattr(curso, "titulo", "") or "").strip(), 9),
        ("CÓDIGO DEL CURSO", (getattr(curso, "codigo", "") or "").strip(), 10),
        ("MATERIAL ENTREGADO", (item_label or "").strip(), 10),
        ("FECHA DE ENTREGA", fecha.strftime("%d/%m/%Y"), 10),
    ]

    for i, (left, right, right_pt) in enumerate(rows):
        c0, c1 = t.rows[i].cells
        c0.text = left
        c1.text = right
        _apply_cell_style(c0, font_pt=10, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
        _apply_cell_style(c1, font_pt=right_pt, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT)

        t.rows[i].height = Cm(0.7)
        t.rows[i].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

    doc.add_paragraph("")

    # ========= Таблица учеников =========
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Table Grid"
    tbl.autofit = False

    # ✅ ВАЖНО: лучше LEFT, иначе Word может "выталкивать" таблицу за край
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    _set_table_indent_zero(tbl)
    _set_table_layout_fixed(tbl)

    # ✅ задаём общую ширину таблицы (иначе Word может масштабировать как хочет)
    _set_table_width_twips(tbl, tw)

    # ✅ считаем ширины колонок (twips)
    w0 = int(tw * 0.07)
    w1 = int(tw * 0.55)
    w2 = int(tw * 0.30)
    w3 = int(tw * 0.08)

    # ✅ применяем ширины к колонкам и всем строкам (включая новые)
    _apply_col_widths_to_all_rows(tbl, [w0, w1, w2, w3])


    hdr = tbl.rows[0].cells
    hdr[0].text = "Nº"
    hdr[1].text = "APELLIDOS Y NOMBRE DEL ALUMN@"
    hdr[2].text = "EMAIL"
    hdr[3].text = "RECIBIDO"

    for c in hdr:
        _apply_cell_style(c, font_pt=10, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    tbl.rows[0].height = Cm(0.7)
    tbl.rows[0].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

    for i, a in enumerate(alumnos_rows, start=1):
        r = tbl.add_row()
        # ✅ Word иногда сбрасывает widths на новых строках
        _apply_col_widths_to_all_rows(tbl, [w0, w1, w2, w3])

        r.height = Cm(0.7)
        r.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

        c0, c1, c2, c3 = r.cells
        c0.text = f"{i}."
        c1.text = (a.get("nombre") or "").strip().upper()
        c2.text = _zwsp_wrap_email(a.get("email") or "")
        c3.text = "Sí" if bool(a.get("received")) else "☐"

        _apply_cell_style(c0, font_pt=10, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER)
        _apply_cell_style(c1, font_pt=10, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT)
        _apply_cell_style(c2, font_pt=10, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT)
        _apply_cell_style(c3, font_pt=10, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.save(str(out_path))

# ------------------------------------------------------------------
# teacher_physical_report_doc — как было, тут менять не обязательно
# ------------------------------------------------------------------
from django.http import FileResponse, Http404
from datetime import date as _date
from pathlib import Path

@login_required
def teacher_physical_report_doc(request, codigo: str):
    curso = get_object_or_404(Curso, codigo=codigo)

    # доступ только teacher этого курса
    is_teacher = Enrol.objects.filter(user=request.user, codigo=curso.codigo, role="teacher").exists()
    if not is_teacher:
        raise Http404

    # какой предмет печатаем
    item_key = (request.GET.get("item_key") or request.GET.get("item") or "").strip()
    if not item_key:
        item_key = "material"  # fallback

    it = CursoPhysicalItem.objects.filter(curso=curso, key=item_key).first()
    item_label = (it.label if it else item_key)

    fecha = _date.today()

    # --- берём студентов курса ---
    enrols = (
        Enrol.objects
        .filter(codigo=curso.codigo, role="student")
        .select_related("user")
        .order_by("created_at", "id")
    )

    user_ids = [e.user_id for e in enrols]

    # --- кто уже подтвердил получение этого предмета ---
    received_ids = set(
        MaterialReceipt.objects
        .filter(curso=curso, item_key=item_key, alumno_id__in=user_ids)
        .values_list("alumno_id", flat=True)
    )

    # --- строим alumnos_rows для DOCX ---
    alumnos_rows = []
    for e in enrols:
        u = e.user

        # если есть UserProfile — берём фамилии оттуда
        profile = getattr(u, "profile", None)
        last1 = (getattr(profile, "last_name1", "") or "").strip()
        last2 = (getattr(profile, "last_name2", "") or "").strip()
        first = (getattr(profile, "first_name", "") or u.first_name or "").strip()

        apellidos = " ".join([x for x in [last1, last2] if x]).strip()
        nombre = f"{apellidos}, {first}".strip().strip(",")
        if not nombre:
            nombre = (u.get_full_name() or u.username or u.email or f"ID {u.pk}").strip()

        alumnos_rows.append({
            "nombre": nombre,
            "email": (u.email or "").strip(),
            "received": (u.pk in received_ids),
        })

    # куда сохраняем
    out_dir = Path(getattr(settings, "MEDIA_ROOT", "/tmp")) / "tmp_docs"
    out_dir.mkdir(parents=True, exist_ok=True)

    # чтобы не затиралось и было понятно что за предмет
    safe_key = "".join(ch for ch in item_key if ch.isalnum() or ch in ("-", "_"))[:40] or "item"
    out_path = out_dir / f"acta_material_{curso.codigo}_{safe_key}.docx"

    build_physical_report_docx(curso, item_label, alumnos_rows, fecha, out_path)
    inject_headers_footers_original(out_path)

    return FileResponse(open(out_path, "rb"), as_attachment=True, filename=out_path.name)
